#!/usr/bin/env python3
"""
Generate three separate 1–2 page client-ready quick-start guides for DensCare.
Each guide is role-specific and can be sent directly to clinic staff.

Outputs:
  DensCare_Administrator_Quick_Start_Guide.docx
  DensCare_Doctor_Quick_Start_Guide.docx
  DensCare_Receptionist_Quick_Start_Guide.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ─────────────────────────────────────────────────
BRAND_DARK   = RGBColor(0x1E, 0x3A, 0x5F)   # deep navy
BRAND_MID    = RGBColor(0x2B, 0x7A, 0x78)   # teal
BRAND_LIGHT  = RGBColor(0xE8, 0xF4, 0xF8)   # light teal bg
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)
GRAY         = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
TABLE_HEADER = RGBColor(0x1E, 0x3A, 0x5F)

# ── Helpers ───────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    tcPr.append(shading)


def _set_cell_borders(cell, color="1E3A5F", sz="4"):
    """Thin border on all sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _remove_table_borders(table):
    """Remove all visible borders from a table (for clean layout)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_paragraph_spacing(para, before=0, after=4, line=None):
    """Set paragraph spacing in points."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before * 20)))   # twips
    spacing.set(qn("w:after"), str(int(after * 20)))
    if line:
        spacing.set(qn("w:line"), str(int(line * 20)))
        spacing.set(qn("w:lineRule"), "auto")


def _make_doc(brand_name="DensCare"):
    """Create a new Document with narrow margins and return it."""
    doc = Document()
    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    return doc


def _add_cover_block(doc, title, subtitle, role_badge, role_color):
    """Add a branded cover block (not a full page — just a compact header)."""
    # Brand name
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=0)
    run = p.add_run("DensCare")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK

    # Subtitle
    p2 = doc.add_paragraph()
    _set_paragraph_spacing(p2, before=0, after=2)
    run2 = p2.add_run("Dental Clinic Management System")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    # Title
    p3 = doc.add_paragraph()
    _set_paragraph_spacing(p3, before=6, after=2)
    run3 = p3.add_run(title)
    run3.font.size = Pt(16)
    run3.font.bold = True
    run3.font.color.rgb = BRAND_DARK

    # Role badge
    p4 = doc.add_paragraph()
    _set_paragraph_spacing(p4, before=0, after=4)
    run4 = p4.add_run(f"  {role_badge}  ")
    run4.font.size = Pt(9)
    run4.font.bold = True
    run4.font.color.rgb = WHITE
    # Badge background
    rPr = run4._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), role_color)
    rPr.append(shd)

    # Thin rule
    p5 = doc.add_paragraph()
    _set_paragraph_spacing(p5, before=0, after=6)
    pPr = p5._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B7A78")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc, text, level=2):
    """Add a coloured section heading."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=10, after=4)
    run = p.add_run(text)
    run.font.size = Pt(12) if level == 2 else Pt(10)
    run.font.bold = True
    run.font.color.rgb = BRAND_MID
    if level == 3:
        run.font.size = Pt(10)
    return p


def _body(doc, text, bold_prefix=None):
    """Add a body paragraph, optionally with a bold prefix."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=3)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.size = Pt(10)
        rb.font.bold = True
        rb.font.color.rgb = BLACK
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    return p


def _bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=2)
    indent = 0.4 + level * 0.4
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(int(indent * 1440)))   # twips (1 inch = 1440)
    ind.set(qn("w:hanging"), str(int(0.2 * 1440)))
    pPr.append(ind)
    # bullet character
    marker = p.add_run("•  ")
    marker.font.size = Pt(10)
    marker.font.color.rgb = BRAND_MID
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.size = Pt(10)
        rb.font.bold = True
        rb.font.color.rgb = BLACK
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    return p


def _numbered_step(doc, number, text):
    """Add a numbered step with bold number."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=2)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(int(0.4 * 1440)))
    ind.set(qn("w:hanging"), str(int(0.3 * 1440)))
    pPr.append(ind)
    rn = p.add_run(f"{number}. ")
    rn.font.size = Pt(10)
    rn.font.bold = True
    rn.font.color.rgb = BRAND_MID
    rt = p.add_run(text)
    rt.font.size = Pt(10)
    rt.font.color.rgb = BLACK
    return p


def _permissions_table(doc, rows_data):
    """Create a compact permissions table with header row."""
    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _remove_table_borders(table)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(11.5)

    # Header
    hdr = table.rows[0]
    for i, txt in enumerate(["What You Can Do", "Details"]):
        cell = hdr.cells[i]
        _set_cell_shading(cell, "1E3A5F")
        _set_cell_borders(cell, "1E3A5F")
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        _set_paragraph_spacing(p, before=2, after=2)

    # Data rows
    for idx, (col1, col2) in enumerate(rows_data):
        row = table.rows[idx + 1]
        bg = "F5F5F5" if idx % 2 == 0 else "FFFFFF"
        for ci, txt in enumerate([col1, col2]):
            cell = row.cells[ci]
            _set_cell_shading(cell, bg)
            _set_cell_borders(cell, "DDDDDD", "2")
            p = cell.paragraphs[0]
            p.text = ""
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
            if ci == 0:
                run.font.bold = True
            _set_paragraph_spacing(p, before=2, after=2)


def _tip_box(doc, text):
    """Add a highlighted tip/note box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "E8F4F8")
    _set_cell_borders(cell, "2B7A78", "6")
    p = cell.paragraphs[0]
    run_icon = p.add_run("💡 Tip:  ")
    run_icon.font.size = Pt(9)
    run_icon.font.bold = True
    run_icon.font.color.rgb = BRAND_MID
    run_text = p.add_run(text)
    run_text.font.size = Pt(9)
    run_text.font.color.rgb = BLACK
    _set_paragraph_spacing(p, before=4, after=4)


def _footer(doc, page_text=""):
    """Add a small footer line."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=8, after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "CCCCCC")
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(f"DensCare — {page_text}   |   For support, contact your system administrator.")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ══════════════════════════════════════════════════════════════════
#  ADMINISTRATOR GUIDE
# ══════════════════════════════════════════════════════════════════

def build_admin_guide():
    doc = _make_doc()
    _add_cover_block(
        doc,
        "Administrator Quick-Start Guide",
        "Step-by-step guide for daily system management",
        "ROLE: ADMINISTRATOR",
        "1E3A5F",
    )

    # ── Signing In ────────────────────────────────────────────────
    _section_heading(doc, "1.  Signing In")
    _numbered_step(doc, 1, "Open DensCare in your web browser.")
    _numbered_step(doc, 2, "Enter your email address and password.")
    _numbered_step(doc, 3, 'Tick "Remember me" if you want DensCare to keep you signed in.')
    _numbered_step(doc, 4, "Click  Sign In.")
    _body(doc, "You will land on the Dashboard showing an overview of your clinic.")

    # ── Dashboard ─────────────────────────────────────────────────
    _section_heading(doc, "2.  Your Dashboard")
    _bullet(doc, "Total patients, today's appointments, active treatment plans, and outstanding invoices.", bold_prefix="Overview Statistics")
    _bullet(doc, "New Patient, New Appointment, New Invoice, New Payment — one-click shortcuts.", bold_prefix="Quick Actions")
    _bullet(doc, "Treatment plans assigned to you, upcoming appointments, and recent activity.", bold_prefix="Widgets")

    # ── Managing Users ────────────────────────────────────────────
    _section_heading(doc, "3.  Managing Users")
    _bullet(doc, "Navigate to  Administration → Users.")
    _bullet(doc, "View all registered users with their name, email, role, and status.")
    _bullet(doc, "Search by name or email. Filter by role or status.")
    _bullet(doc, "Click a user to view their profile, change their role, or activate/deactivate them.")
    _bullet(doc, "Click  Add User  to create a new staff account. Fill in name, email, and password, then assign a role.")
    _tip_box(doc, "You cannot deactivate your own account or remove the last administrator from the system.")

    # ── Approving New Registrations ───────────────────────────────
    _section_heading(doc, "4.  Approving New Registrations")
    _bullet(doc, "Navigate to  Administration → Pending Approvals.")
    _bullet(doc, "New users who register appear here waiting for your approval.")
    _bullet(doc, "Select a role for the user from the dropdown (e.g., Doctor, Receptionist).")
    _bullet(doc, "Click  Approve  to activate the account, or  Reject  to deny access.")

    # ── Managing Doctors ──────────────────────────────────────────
    _section_heading(doc, "5.  Managing Doctors")
    _bullet(doc, "Navigate to  Clinical → Doctors.")
    _bullet(doc, "Click  Add Doctor  to register a new doctor profile.")
    _bullet(doc, "Fill in: name, registration number, specialization, consultation fee, and duration.")
    _bullet(doc, "Link the doctor profile to an existing user account (must have a doctor role).")
    _bullet(doc, "Manage each doctor's weekly schedule (available days and time slots).")
    _bullet(doc, "Mark doctors as Active/Inactive or Available/Unavailable for appointments.")

    # ── Patient & Appointment Management ──────────────────────────
    _section_heading(doc, "6.  Patients & Appointments")
    _bullet(doc, "You can register new patients, edit patient details, and deactivate patient accounts.")
    _bullet(doc, "Schedule, reschedule, or cancel any appointment for any doctor.")
    _bullet(doc, "View and manage all patient records, treatment plans, and prescriptions.")

    # ── Billing ───────────────────────────────────────────────────
    _section_heading(doc, "7.  Billing & Financial")
    _bullet(doc, "Create, edit, issue, and delete draft invoices.")
    _bullet(doc, "Record, complete, void, and delete payments.")
    _bullet(doc, "Generate, view, and regenerate receipts.")
    _bullet(doc, "Process refunds (approve, reject, complete) and manage credit notes (issue, void, apply).")
    _bullet(doc, "Access the Billing Dashboard for financial overview and reports.")

    # ── Permissions table ─────────────────────────────────────────
    _section_heading(doc, "8.  What You Can and Cannot Do")

    _permissions_table(doc, [
        ("✅  Create / edit / deactivate users",          "Full user management"),
        ("✅  Approve or reject new registrations",       "Administration → Pending Approvals"),
        ("✅  Register and manage doctors",               "Full doctor profile and schedule control"),
        ("✅  Register and manage patients",              "Create, edit, deactivate patients"),
        ("✅  Manage all appointments",                   "Schedule, reschedule, cancel for any doctor"),
        ("✅  View and manage all patient records",       "Clinical records, diagnoses, prescriptions"),
        ("✅  Full billing access",                       "Invoices, payments, receipts, refunds, credit notes"),
        ("✅  Delete draft invoices and pending payments","Only administrator can delete"),
        ("✅  View audit logs for patient records",       "Administration-level compliance review"),
        ("❌  Cannot deactivate your own account",        "System safety protection"),
        ("❌  Cannot remove the last admin role",         "Prevents system lockout"),
    ])

    _footer(doc, "Administrator Quick-Start Guide  •  Version 1.0")
    doc.save("DensCare_Administrator_Quick_Start_Guide.docx")
    print("[OK] DensCare_Administrator_Quick_Start_Guide.docx")


# ══════════════════════════════════════════════════════════════════
#  DOCTOR GUIDE
# ══════════════════════════════════════════════════════════════════

def build_doctor_guide():
    doc = _make_doc()
    _add_cover_block(
        doc,
        "Doctor Quick-Start Guide",
        "Step-by-step guide for clinical daily workflow",
        "ROLE: DOCTOR",
        "2B7A78",
    )

    # ── Signing In ────────────────────────────────────────────────
    _section_heading(doc, "1.  Signing In")
    _numbered_step(doc, 1, "Open DensCare in your web browser.")
    _numbered_step(doc, 2, "Enter your email address and password.")
    _numbered_step(doc, 3, 'Tick "Remember me" to stay signed in between sessions.')
    _numbered_step(doc, 4, "Click  Sign In.")
    _body(doc, "You will land on the Dashboard with your appointments and treatment plans.")

    # ── Dashboard ─────────────────────────────────────────────────
    _section_heading(doc, "2.  Your Dashboard")
    _bullet(doc, "See your total patients, today's appointments, active treatment plans, and outstanding invoices at a glance.", bold_prefix="Overview Statistics")
    _bullet(doc, "Quick Actions for creating new patients, appointments, invoices, and payments.", bold_prefix="Quick Actions")
    _bullet(doc, "My Treatment Plans shows plans assigned to you. Upcoming Appointments lists what's ahead.", bold_prefix="Widgets")

    # ── Working with Appointments ─────────────────────────────────
    _section_heading(doc, "3.  Your Appointments")
    _bullet(doc, "Navigate to  Clinical → Appointments  to see all appointments.")
    _bullet(doc, "Use the date filter to view today's, this week's, or this month's schedule.")
    _bullet(doc, "Click on an appointment to see patient details, status, and notes.")
    _bullet(doc, "You can reschedule or cancel appointments as needed.")

    # ── Patient Records ───────────────────────────────────────────
    _section_heading(doc, "4.  Patient Records (Clinical Work)")
    _body(doc, "This is where you do most of your clinical documentation.")
    _numbered_step(doc, 1, "Navigate to  Clinical → Patient Records  and open a record.")
    _numbered_step(doc, 2, "The  Clinical Info  tab is where you enter chief complaint, clinical notes, doctor remarks, allergies, medical history, and dental history.")
    _numbered_step(doc, 3, "Switch to the  Prescriptions  tab to create and manage prescriptions for this record.")
    _numbered_step(doc, 4, "Use the  Follow-ups  tab to schedule follow-up visits.")
    _numbered_step(doc, 5, "Use the  Attachments  tab to upload files (X-rays, consent forms, etc.).")
    _tip_box(doc, "When a record is finalised, it becomes read-only for medico-legal compliance. Make sure all information is complete before finalising.")

    # ── Diagnoses ─────────────────────────────────────────────────
    _section_heading(doc, "5.  Adding a Diagnosis")
    _bullet(doc, "Within a patient record, go to the  Clinical Info  tab.")
    _bullet(doc, "Add diagnoses with ICD codes, description, and status (Active, Resolved, etc.).")
    _bullet(doc, "Multiple diagnoses can be recorded per visit.")

    # ── Prescriptions ─────────────────────────────────────────────
    _section_heading(doc, "6.  Creating a Prescription")
    _numbered_step(doc, 1, "Open the patient record →  Prescriptions  tab.")
    _numbered_step(doc, 2, "Click  New Prescription.")
    _numbered_step(doc, 3, "Add medicines with: medicine name, dosage, frequency, duration, and instructions.")
    _numbered_step(doc, 4, "Save the prescription.")
    _numbered_step(doc, 5, "Click the print icon to print, or use the browser's Print → Save as PDF to download.")
    _tip_box(doc, "To download a prescription as PDF, click the print icon, then in the print dialog choose 'Save as PDF' as the destination.")

    # ── Treatment Plans ───────────────────────────────────────────
    _section_heading(doc, "7.  Treatment Plans")
    _bullet(doc, "Navigate to  Clinical → Treatment Plans  to view all plans.")
    _bullet(doc, "Click  New Plan  to create a treatment plan for a patient.")
    _bullet(doc, "Add procedures from the procedure catalogue, set status, and assign to a doctor.")
    _bullet(doc, "Track plan status: Draft → Submitted → Proposed → Accepted → In Progress → Completed.")

    # ── Attachments ───────────────────────────────────────────────
    _section_heading(doc, "8.  File Attachments")
    _bullet(doc, "Upload images, PDFs, and documents from the  Attachments  tab in a patient record.")
    _bullet(doc, "Supported formats: JPEG, PNG, GIF, PDF, DOC, DOCX. Maximum 10 MB per file.")
    _bullet(doc, "Download or delete attachments as needed.")

    # ── Billing (read + create) ──────────────────────────────────
    _section_heading(doc, "9.  Billing")
    _bullet(doc, "View the Billing Dashboard for financial overview.", bold_prefix="View:")
    _bullet(doc, "Create invoices, record payments, generate receipts.", bold_prefix="Create:")
    _bullet(doc, "Issue or cancel invoices, complete or void payments, process refunds.", bold_prefix="Manage:")
    _bullet(doc, "Only administrators can delete draft invoices or pending payments.", bold_prefix="Restriction:")

    # ── Permissions table ─────────────────────────────────────────
    _section_heading(doc, "10.  What You Can and Cannot Do")

    _permissions_table(doc, [
        ("✅  View your appointments and schedule",         "Clinical → Appointments"),
        ("✅  View and edit patient records",               "Clinical records, diagnoses, prescriptions"),
        ("✅  Create and manage prescriptions",             "Add medicines, print/download PDF"),
        ("✅  Create and manage treatment plans",           "Add procedures, track status"),
        ("✅  Upload and manage attachments",               "X-rays, consent forms, documents"),
        ("✅  Create invoices and record payments",         "Financial → Invoices / Payments"),
        ("✅  Generate and view receipts",                  "Automatic after payment completion"),
        ("✅  Process refunds and credit notes",            "Create requests, approve/reject"),
        ("❌  Cannot register or deactivate patients",      "Admin and Receptionist only"),
        ("❌  Cannot manage user accounts or roles",        "Administrator only"),
        ("❌  Cannot manage doctor profiles or schedules",  "Administrator only"),
        ("❌  Cannot delete invoices or payments",          "Administrator only"),
    ])

    _footer(doc, "Doctor Quick-Start Guide  •  Version 1.0")
    doc.save("DensCare_Doctor_Quick_Start_Guide.docx")
    print("[OK] DensCare_Doctor_Quick_Start_Guide.docx")


# ══════════════════════════════════════════════════════════════════
#  RECEPTIONIST GUIDE
# ══════════════════════════════════════════════════════════════════

def build_receptionist_guide():
    doc = _make_doc()
    _add_cover_block(
        doc,
        "Receptionist Quick-Start Guide",
        "Step-by-step guide for front-desk daily workflow",
        "ROLE: RECEPTIONIST",
        "E07A5F",
    )

    # ── Signing In ────────────────────────────────────────────────
    _section_heading(doc, "1.  Signing In")
    _numbered_step(doc, 1, "Open DensCare in your web browser.")
    _numbered_step(doc, 2, "Enter your email address and password.")
    _numbered_step(doc, 3, 'Tick "Remember me" to stay signed in between sessions.')
    _numbered_step(doc, 4, "Click  Sign In.")
    _body(doc, "You will land on the Dashboard showing today's overview.")

    # ── Dashboard ─────────────────────────────────────────────────
    _section_heading(doc, "2.  Your Dashboard")
    _bullet(doc, "Total patients, today's appointments, active treatment plans, and outstanding invoices.", bold_prefix="Overview Statistics")
    _bullet(doc, "Quick shortcuts to register a patient, schedule an appointment, create an invoice, or record a payment.", bold_prefix="Quick Actions")

    # ── Registering a Patient ─────────────────────────────────────
    _section_heading(doc, "3.  Registering a New Patient")
    _numbered_step(doc, 1, 'Navigate to  Clinical → Patients  and click  New Patient.')
    _numbered_step(doc, 2, "Fill in the required fields: full name, date of birth, phone number, and gender.")
    _numbered_step(doc, 3, "Optionally add: email, address, emergency contact, blood group, and medical alerts.")
    _numbered_step(doc, 4, "Click  Save  to create the patient record.")
    _tip_box(doc, "Use the search bar on the Patients page to quickly find existing patients by name or phone number.")

    # ── Searching for a Patient ───────────────────────────────────
    _section_heading(doc, "4.  Finding a Patient")
    _bullet(doc, "Go to  Clinical → Patients.")
    _bullet(doc, "Type the patient's name, phone number, or email in the search bar.")
    _bullet(doc, "Use the Status filter (Active / Inactive) to narrow results.")
    _bullet(doc, "Click on a patient row to open their full profile.")

    # ── Scheduling Appointments ───────────────────────────────────
    _section_heading(doc, "5.  Scheduling an Appointment")
    _numbered_step(doc, 1, "Navigate to  Clinical → Appointments  and click  New Appointment.")
    _numbered_step(doc, 2, "Select the patient from the dropdown (search by name).")
    _numbered_step(doc, 3, "Select the doctor from the dropdown.")
    _numbered_step(doc, 4, "Choose the appointment date and time.")
    _numbered_step(doc, 5, "Add any notes if needed, then click  Save.")
    _body(doc, "The appointment appears in the doctor's schedule and the clinic calendar.")

    # ── Rescheduling / Cancelling ─────────────────────────────────
    _section_heading(doc, "6.  Rescheduling or Cancelling an Appointment")
    _bullet(doc, "Open the appointment from the Appointments list.")
    _bullet(doc, "To reschedule: change the date/time and save.")
    _bullet(doc, "To cancel: click the Cancel button and confirm.")
    _tip_box(doc, "Cancelled appointments are kept in the system for record-keeping but will no longer appear in the active schedule.")

    # ── Patient Records ───────────────────────────────────────────
    _section_heading(doc, "7.  Patient Records")
    _bullet(doc, "You can view patient records and add clinical notes, prescriptions, and attachments.")
    _bullet(doc, "Navigate to  Clinical → Patient Records  and open a record.")
    _bullet(doc, "Switch between tabs: Clinical Info, Prescriptions, Follow-ups, and Attachments.")
    _bullet(doc, "You cannot finalise or delete records — that is restricted to doctors and administrators.")

    # ── Billing ───────────────────────────────────────────────────
    _section_heading(doc, "8.  Billing & Payments")
    _numbered_step(doc, 1, "Navigate to  Financial → Invoices  to view or create invoices.")
    _numbered_step(doc, 2, "Click  New Invoice  → select patient → add line items (procedures/fees) → Save as Draft.")
    _numbered_step(doc, 3, "Click  Issue  to finalise the invoice (it becomes read-only).")
    _numbered_step(doc, 4, "Navigate to  Financial → Payments  and click  New Payment.")
    _numbered_step(doc, 5, "Select the invoice, enter payment amount and method (Cash, Card, Bank Transfer, etc.).")
    _numbered_step(doc, 6, "Click  Complete  to finalise the payment. A receipt is generated automatically.")
    _bullet(doc, "You can also process refunds and credit notes.", bold_prefix="Refunds & Credit Notes:")

    # ── Permissions table ─────────────────────────────────────────
    _section_heading(doc, "9.  What You Can and Cannot Do")

    _permissions_table(doc, [
        ("✅  Register new patients",                       "Clinical → Patients → New Patient"),
        ("✅  Edit patient details",                        "Update contact info, medical alerts, etc."),
        ("✅  Schedule, reschedule, cancel appointments",   "Clinical → Appointments"),
        ("✅  View and edit patient records",               "Clinical records, prescriptions, notes"),
        ("✅  Upload file attachments",                     "X-rays, consent forms, documents"),
        ("✅  Create and issue invoices",                   "Financial → Invoices"),
        ("✅  Record and complete payments",                "Financial → Payments"),
        ("✅  Generate and view receipts",                  "Automatic after payment completion"),
        ("✅  Process refunds and credit notes",            "Create, approve, reject, complete"),
        ("❌  Cannot deactivate patients",                  "Administrator only"),
        ("❌  Cannot finalise or delete patient records",   "Doctors and Administrators only"),
        ("❌  Cannot manage users, doctors, or roles",      "Administrator only"),
        ("❌  Cannot delete invoices or payments",          "Administrator only"),
    ])

    _footer(doc, "Receptionist Quick-Start Guide  •  Version 1.0")
    doc.save("DensCare_Receptionist_Quick_Start_Guide.docx")
    print("[OK] DensCare_Receptionist_Quick_Start_Guide.docx")


# ══════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    build_admin_guide()
    build_doctor_guide()
    build_receptionist_guide()
    print("\nAll three role guides generated successfully.")
