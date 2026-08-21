"""
DensCare — Dental Clinic Management System
User Manual & Client Training Guide (.docx generator)

Generates `DensCare_User_Manual_Client_Training_Guide.docx` from the audited
DensCare implementation (backend RBAC + frontend UI). Run:

    python generate_denscare_user_manual.py

Requires: python-docx (pip install python-docx)
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─────────────────────────────────────────────────────────────────────
# Brand palette (mirrors the DensCare frontend design tokens)
# ─────────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1D, 0x4E, 0xD8)        # primary-700
NAVY_DARK = RGBColor(0x1E, 0x29, 0x3B)   # neutral-900
BLUE = RGBColor(0x25, 0x63, 0xEB)        # primary-600
SLATE = RGBColor(0x47, 0x55, 0x69)       # neutral-600
LIGHT_BLUE_FILL = "EFF6FF"               # primary-50
HEADER_FILL = "1E3A8A"                   # deep navy table header
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SUCCESS = RGBColor(0x05, 0x96, 0x69)
WARNING = RGBColor(0xD9, 0x77, 0x06)
DANGER = RGBColor(0xDC, 0x26, 0x26)

BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"

A4_W = Cm(21.0)
A4_H = Cm(29.7)
MARGIN = Cm(2.0)

OUT_PATH = "DensCare_User_Manual_Client_Training_Guide.docx"
NAME_PLATE = os.path.join("frontend", "src", "assets", "images", "name.png")
LOGO = os.path.join("frontend", "src", "assets", "images", "logo.png")

CHECK = "\u2713"      # ✓
CROSS = "\u2014"      # —
ARROW = "\u2192"      # →

# ─────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────

def rgb_hex(color: RGBColor) -> str:
    """Return an RRGGBB hex string for an RGBColor (e.g. 1D4ED8)."""
    return "%02X%02X%02X" % (color[0], color[1], color[2])


def insert_ordered(parent, new_el, order: list[str]) -> None:
    """Insert an OOXML child element respecting the schema child order.

    Word can flag out-of-order elements inside `w:tcPr` / `w:tblPr` /
    `w:settings`, so elements appended with a raw `append()` are placed
    before the first existing child that must follow them.
    """
    tag = new_el.tag.split("}")[-1]
    idx = order.index(tag)
    for child in parent:
        ctag = child.tag.split("}")[-1]
        if ctag in order and order.index(ctag) > idx:
            child.addprevious(new_el)
            return
    parent.append(new_el)


TC_PR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders",
    "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign",
    "hideMark",
]

TBL_PR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
    "tblCaption", "tblDescription",
]

# Paragraph property child order (subset) so w:pBdr lands before spacing/jc.
P_PR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
    "sectPr", "pPrChange",
]

# Subset of the CT_Settings child order that can appear in the default
# python-docx template, around where `w:updateFields` belongs.
SETTINGS_ORDER = [
    "updateFields", "hdrShapeDefaults", "footnotePr", "endnotePr",
    "compat", "rsids", "mathPr", "uiCompat97To2003", "attachedSchema",
    "themeFontLang", "clrSchemeMapping", "doNotIncludeSubdocsInStats",
    "doNotAutoCompressPictures", "forceUpgrade", "captions",
    "readModeInkLockDown", "smartTagType", "decimalSymbol", "listSeparator",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    insert_ordered(tc_pr, shd, TC_PR_ORDER)


def set_repeat_table_header(row) -> None:
    """Repeat the header row across printed pages."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_box(paragraph, color="C7D2FE", sz=6) -> None:
    """Draw a thin rounded-less box around a paragraph (used for screenshots)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    insert_ordered(p_pr, p_bdr, P_PR_ORDER)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, t, fld_end):
        run._r.append(el)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    insert_ordered(settings, update, SETTINGS_ORDER)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Heading 1", 17, NAVY_DARK, 14, 6),
        ("Heading 2", 13.5, NAVY, 10, 4),
        ("Heading 3", 11.5, BLUE, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        # Ensure headings use a latin font too (Word quirk with Calibri)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), HEADING_FONT)
        rfonts.set(qn("w:hAnsi"), HEADING_FONT)

    # Title style for the cover
    title = doc.styles["Title"]
    title.font.name = HEADING_FONT
    title.font.size = Pt(34)
    title.font.bold = True
    title.font.color.rgb = NAVY_DARK

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = HEADING_FONT
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = BLUE


def para(doc, text, bold=False, italic=False, size=None, color=None,
         align=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def rich_para(doc, segments, align=None, space_after=None):
    """segments: list of (text, bold, italic, color) tuples."""
    p = doc.add_paragraph()
    for text, bold, italic, color in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(1.1)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def steps(doc, items, start=1):
    """Numbered procedure steps with hanging indent (stable numbering)."""
    for i, text in enumerate(items, start=start):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{i}.  ")
        run.bold = True
        run.font.color.rgb = BLUE
        if isinstance(text, tuple):
            lead, rest = text
            r1 = p.add_run(lead)
            r1.bold = True
            p.add_run(rest)
        else:
            p.add_run(text)


def add_table(doc, headers, rows, col_widths=None, header_fill=HEADER_FILL,
              zebra=True, font_size=9.5, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, header_fill)
        p.paragraph_format.space_after = Pt(1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                txt, is_bold = val
            else:
                txt, is_bold = val, False
            run = p.add_run(str(txt))
            run.font.size = Pt(font_size)
            run.bold = is_bold or (first_col_bold and j == 0)
            p.paragraph_format.space_after = Pt(1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if zebra and i % 2 == 1:
                set_cell_shading(cell, "F1F5F9")
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = w
    return table


SCREENSHOT_DIR = os.path.join("docs", "manual-screenshots")
_FIG = {"n": 0}


def screenshot(doc, filename, caption, width_cm=14.5):
    """Embed a real captured screenshot with a thin frame and a numbered caption.

    Skips gracefully (with a visible note) if the PNG is not present, so the
    manual still builds on machines without the screenshot files.
    """
    path = os.path.join(SCREENSHOT_DIR, filename)
    if not os.path.exists(path):
        para(doc, f"[Screenshot not available: {filename}]", italic=True,
             color=DANGER, space_after=8)
        return None
    _FIG["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    set_paragraph_box(p)
    para(doc, f"Figure {_FIG['n']} — {caption}", italic=True, size=9,
         color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    return p


def callout(doc, title, text, fill=LIGHT_BLUE_FILL, border=NAVY, title_color=None):
    """Single-cell callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p0 = cell.paragraphs[0]
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = title_color or border
    p0.paragraph_format.space_after = Pt(2)
    p1 = cell.add_paragraph()
    r1 = p1.add_run(text)
    r1.font.size = Pt(9.5)
    p1.paragraph_format.space_after = Pt(1)
    # Border colour (inserted in schema order within tblPr)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), rgb_hex(border) if isinstance(border, RGBColor) else "2563EB")
        borders.append(el)
    insert_ordered(tbl_pr, borders, TBL_PR_ORDER)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_footer(doc, section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DensCare — User Manual & Client Training Guide  |  Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = SLATE
    add_field(p, "PAGE")
    for r in p.runs:
        r.font.size = Pt(8.5)
        r.font.color.rgb = SLATE


def build_header(doc, section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("DensCare — Dental Clinic Management System")
    run.font.size = Pt(8.5)
    run.font.color.rgb = SLATE
    run.italic = True


def add_role_badge(doc, role_name):
    return rich_para(
        doc,
        [(role_name, True, False, NAVY)],
        space_after=2,
    )


# ─────────────────────────────────────────────────────────────────────
# Document construction
# ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page geometry (A4)
    for section in doc.sections:
        section.page_width = A4_W
        section.page_height = A4_H
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)

    setup_styles(doc)
    enable_update_fields(doc)

    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True

    # ════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    if os.path.exists(NAME_PLATE):
        doc.add_picture(NAME_PLATE, width=Cm(12.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("DensCare")
    run.font.name = HEADING_FONT
    run.font.size = Pt(44)
    run.font.bold = True
    run.font.color.rgb = NAVY_DARK

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub1.add_run("Dental Clinic Management System")
    run.font.size = Pt(20)
    run.font.color.rgb = BLUE
    run.bold = True

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("User Manual & Client Training Guide")
    run.font.size = Pt(16)
    run.font.color.rgb = SLATE

    doc.add_paragraph()
    doc.add_paragraph()

    cover_table = doc.add_table(rows=4, cols=2)
    cover_table.style = "Table Grid"
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("Version", "1.1"),
        ("Document date", "12 August 2026"),
        ("Prepared for", "Dental Clinic Client"),
        ("Prepared by", "DensCare Project Team"),
    ]
    for i, (k, v) in enumerate(meta):
        c0, c1 = cover_table.rows[i].cells
        c0.text = ""
        p = c0.paragraphs[0]
        r = p.add_run(k)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY
        set_cell_shading(c0, LIGHT_BLUE_FILL)
        c1.text = ""
        p = c1.paragraphs[0]
        r = p.add_run(v)
        r.font.size = Pt(10.5)
    for row in cover_table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(9.5)

    doc.add_paragraph()
    doc.add_paragraph()
    if os.path.exists(LOGO):
        doc.add_picture(LOGO, width=Cm(2.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u',
              "Table of contents will appear here. In Microsoft Word, right-click "
              "and choose \u201cUpdate Field\u201d (or press F9) to generate it. "
              "In LibreOffice, right-click and choose \u201cUpdate Index\u201d.")
    para(doc, "", space_after=2)
    callout(
        doc,
        "Note",
        "This manual was written for the current version of DensCare and matches "
        "the application exactly as delivered. Every procedure, permission and "
        "status described here reflects what the system actually does today.",
    )
    page_break(doc)

    # ── Section 1 ───────────────────────────────────────────────────
    doc.add_heading("1.  Introduction", level=1)

    doc.add_heading("1.1  What is DensCare?", level=2)
    para(doc,
         "DensCare is a web-based Dental Clinic Management System that helps a "
         "dental clinic run its day-to-day operations from one place. It keeps all "
         "patient information, appointments, clinical records, treatment plans, "
         "prescriptions and billing in a single, organised system that authorised "
         "staff can access from a web browser on a computer, tablet or phone.")
    para(doc,
         "DensCare replaces paper registers, loose files and separate billing "
         "notebooks with one central system. When a patient visits, the clinic can "
         "register them, book their appointment, record their treatment, issue a "
         "prescription and bill them — all using the same patient record.")

    doc.add_heading("1.2  What can DensCare do?", level=2)
    for b in (
        ("Centralised patient management", " — keep one master record for every patient, including contact details, emergency contact and remarks."),
        ("Appointment management", " — schedule, view, reschedule and cancel patient appointments with the clinic's doctors."),
        ("Clinical records", " — record each consultation: complaints, clinical notes, diagnoses, medical history and more."),
        ("Treatment planning", " — create structured treatment plans made up of dental procedures, track progress and get patient approval."),
        ("Prescriptions", " — write prescriptions with medicines, dosage and instructions, and print or save them as PDF."),
        ("Billing", " — create invoices, record payments, generate receipts, handle refunds and credit notes."),
        ("User management and access control", " — each staff member signs in with their own account and role, and the system controls what they can see and do."),
        ("Dashboard visibility", " — a single home screen shows quick actions, upcoming appointments and active treatment plans."),
    ):
        bullet(doc, b[1], bold_prefix=b[0])

    doc.add_heading("1.3  Who is this manual for?", level=2)
    para(doc,
         "This manual is written for the clinic staff who will use DensCare every "
         "day: administrators, doctors, receptionists and dental assistants. It is "
         "not a technical document — it explains how to operate the system in plain "
         "language. Each role has its own guide, so staff can focus on the parts of "
         "the system they actually use.")

    doc.add_heading("1.4  How to use this manual", level=2)
    bullet(doc, "Chapters 2 and 3 give an overview of how the system works and the staff roles.")
    bullet(doc, "Chapters 4 to 7 are separate guides for each role — start with the guide for your role.")
    bullet(doc, "Chapters 8 to 19 explain every module and screen in detail.")
    bullet(doc, "Chapter 20 contains complete step-by-step scenarios for common daily situations.")
    bullet(doc, "Chapters 21 to 25 are quick references: FAQs, permissions, training checklists and troubleshooting.")
    bullet(doc, "Chapter 26 lists features that are planned but not yet available.")

    page_break(doc)

    # ── Section 2 ───────────────────────────────────────────────────
    doc.add_heading("2.  How DensCare Works", level=1)

    doc.add_heading("2.1  The clinic workflow in one picture", level=2)
    para(doc,
         "The flow below shows a typical patient journey through the clinic and "
         "how the different parts of DensCare work together. Not every visit uses "
         "every step — for example, a follow-up visit may skip registration, and "
         "an over-the-counter payment may not need a new clinical record.")

    flow = [
        ("1. Patient Registration", "The receptionist or an administrator creates a patient record (or finds an existing one)."),
        ("2. Appointment", "An appointment is booked with a doctor and a suitable date and time."),
        ("3. Clinical Consultation", "The doctor opens the appointment and reviews the patient's history."),
        ("4. Diagnosis / Clinical Record", "The doctor records the consultation: complaints, findings, diagnoses and medical history."),
        ("5. Treatment Plan", "For planned work, the doctor creates a treatment plan with the procedures involved and obtains patient acceptance."),
        ("6. Prescription", "The doctor writes a prescription with the medicines, dosage and instructions, and can print it."),
        ("7. Invoice", "A bill (invoice) is created for the treatment or consultation and issued."),
        ("8. Payment", "The patient pays. Staff record the payment (cash, card, UPI, etc.)."),
        ("9. Receipt / Settlement", "A receipt is generated for the completed payment and can be printed for the patient."),
    ]
    for name, desc in flow:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ARROW + "  ")
        run.font.color.rgb = BLUE
        run.bold = True
        r1 = p.add_run(name + " — ")
        r1.bold = True
        p.add_run(desc)

    doc.add_heading("2.2  How the modules connect", level=2)
    para(doc,
         "DensCare is not a set of separate screens — the modules share data so "
         "information only has to be entered once:")
    bullet(doc, "A patient record is the anchor. Appointments, clinical records, treatment plans, invoices and payments all refer back to one patient.")
    bullet(doc, "An appointment can lead to a clinical record, which holds diagnoses, prescriptions, follow-ups and attachments for that visit.")
    bullet(doc, "A treatment plan is built from the procedure catalogue and can be linked to an invoice so planned procedures become billable items.")
    bullet(doc, "An invoice is created from services/items, issued, and then paid. Payments are allocated to invoices, receipts are generated from completed payments, and refunds or credit notes adjust the financial picture.")
    bullet(doc, "The billing dashboard summarises the financial totals and recent activity, so staff always know the clinic's overall position.")

    callout(doc, "Important",
            "DensCare controls what each staff member can see and do based on their "
            "role. The sections below explain each role and its permissions.")

    page_break(doc)

    # ── Section 3 ───────────────────────────────────────────────────
    doc.add_heading("3.  User Roles", level=1)

    para(doc,
         "DensCare uses roles to decide what each person can access. Every staff "
         "member signs in with their own account, and the account carries one role. "
         "The role controls which modules appear in the menu and — more importantly "
         "— what actions the system allows. There are seven roles in the delivered "
         "system.")

    doc.add_heading("3.1  The roles at a glance", level=2)
    add_table(
        doc,
        ["Role", "Who it is", "Main use of DensCare"],
        [
            [("Administrator", True), "The clinic's system administrator / manager.", "Manages staff accounts, approvals, doctors, patients and the full billing cycle."],
            [("Chief Doctor", True), "The senior/head doctor of the clinic.", "Runs the clinic clinically, manages doctors, and has administrator-level access to most management screens."],
            [("General Doctor", True), "A regular dentist.", "Sees patients, records consultations, writes prescriptions, creates treatment plans and bills."],
            [("Specialist Doctor", True), "A dentist specialising in a particular field.", "Same clinical day-to-day access as a general doctor."],
            [("Consulting Doctor", True), "A doctor who consults on specific cases.", "Same clinical day-to-day access as a general doctor."],
            [("Receptionist", True), "The front-desk staff.", "Registers patients, manages appointments and handles billing at the counter."],
            [("Dental Assistant", True), "A chair-side assistant.", "Supports billing activity (creating draft invoices, payments, refunds, credit notes and receipts) but does not have clinical or workflow access."],
        ],
        col_widths=[Cm(3.4), Cm(5.2), Cm(7.4)],
    )

    doc.add_heading("3.2  Doctor roles", level=2)
    para(doc,
         "General Doctor, Specialist Doctor, Consulting Doctor and Chief Doctor are "
         "all doctor-level roles. They share the same clinical permissions — the "
         "distinction is professional (which doctor is assigned to which case), not "
         "a difference in system access. The Chief Doctor is special: the system "
         "also treats the Chief Doctor as an administrator for management screens "
         "(user approvals, doctor management and other administrator functions).")

    doc.add_heading("3.3  How roles affect what you see", level=2)
    para(doc,
         "Most of the menu is the same for every role. Two menu items are only "
         "visible to administrator-level accounts: Users and Pending Approvals. For "
         "everything else, the menu may look the same, but the system itself "
         "enforces permissions — if you open a screen your role cannot use, "
         "DensCare shows a clear \u201cyou do not have permission\u201d message instead of "
         "the data. Chapter 22 contains the complete permission table.")

    page_break(doc)

    # ── Section 4 ───────────────────────────────────────────────────
    doc.add_heading("4.  Administrator User Guide", level=1)

    doc.add_heading("4.1  Signing in", level=2)
    steps(doc, [
        "Open DensCare in your web browser using the clinic's application address.",
        "Enter your email address in the Email field.",
        "Enter your password in the Password field.",
        ("Tick \u201cRemember me\u201d", " if you are using your own computer and want DensCare to keep you signed in."),
        "Click the Sign In button.",
        ("You arrive at the Dashboard", " — the home screen described in Chapter 17."),
    ])
    callout(doc, "Cannot sign in?",
            "Check that the account has been approved and activated by an "
            "administrator (see Chapter 25, Troubleshooting).")

    screenshot(doc, "01-login.png",
               "The Sign-in screen — enter your email and password, or request a new account.")

    doc.add_heading("4.2  The Dashboard", level=2)
    para(doc,
         "After signing in you see the Dashboard. As an administrator you use it as "
         "a starting point every day: it shows Quick Actions (New Patient, Schedule "
         "Appointment, Create Invoice, View Appointments), your active treatment "
         "plans and today's appointments. The Overview statistics cards and Recent "
         "Activity list show sample figures in the current version — the live "
         "financial numbers live in the Billing Dashboard (Chapter 14).")

    screenshot(doc, "02-dashboard.png",
               "The Dashboard after signing in as an administrator — quick actions, treatment plans and today's appointments.")

    doc.add_heading("4.3  Managing users", level=2)
    para(doc, "Only administrator-level accounts (Administrator and Chief Doctor) can manage users.")
    steps(doc, [
        "Open Users from the Administration group in the left menu.",
        ("View the list of staff accounts", " — you can search by name or email, and filter by role and status (Pending / Active / Inactive)."),
        ("Add a new user directly", " — click Add User on the Users screen, enter the full name, email and password, choose the role, and save. DensCare registers the account and approves it with the chosen role in the same flow. If automatic approval cannot complete, the account stays in the pending queue for approval (section 4.4)."),
        ("Open a user", " to see their profile: full name, email, assigned role, status and account dates."),
        ("Change a user's role", " — on the user's profile, change the role and save. The system prevents you from changing your own role."),
        ("Activate or deactivate an account", " — an inactive account can no longer sign in. The system prevents you from deactivating your own account."),
    ])

    screenshot(doc, "19-users.png",
               "The Users screen — staff accounts with search and filters.")

    doc.add_heading("4.4  Approving new registrations", level=2)
    para(doc,
         "New staff can request an account themselves using the \u201cCreate an "
         "account\u201d link on the sign-in screen. Their request arrives as a pending "
         "registration and can only sign in after an administrator approves it.")
    steps(doc, [
        "Open Pending Approvals from the Administration group in the left menu.",
        "Review the list of staff members waiting for approval (name, email, requested date).",
        ("Approve a person", " — choose the role they should have (Administrator, Chief Doctor, General Doctor, Specialist Doctor, Consulting Doctor, Receptionist or Dental Assistant) and confirm. The account becomes Active and they can sign in."),
        ("Alternatively, reject/deactivate a request", " — the person cannot sign in."),
    ])

    screenshot(doc, "20-pending-approvals.png",
               "Pending Approvals — new staff registrations waiting for an administrator.")

    doc.add_heading("4.5  Managing doctors", level=2)
    para(doc,
         "A doctor profile is linked to a user account. Only administrator-level "
         "accounts can create and manage doctor profiles, specializations and "
         "schedules.")
    steps(doc, [
        "Open Doctors from the Clinical group in the left menu.",
        ("Create a doctor profile", " — pick the user account for that doctor (a doctor role account), then enter their details: date of birth, gender, primary phone, address, qualification, registration number, years of experience, consultation fee and duration, languages, biography and emergency contact."),
        ("Assign specializations", " — from the doctor's profile, assign one or more specializations from the master list and mark the primary one."),
        ("Set the weekly schedule", " — add the doctor's working sessions (day of the week with start and end time) and save."),
        ("Manage availability", " — use the Available / Not Available toggle to control whether the doctor can take new appointments, and the On Leave toggle when they are away."),
        ("Deactivate a doctor profile", " when a doctor leaves the clinic."),
    ])
    para(doc,
         "The specializations themselves (the master list, e.g. Endodontics, "
         "Orthodontics) are managed from the same area by an administrator, and the "
         "Procedure Catalogue is managed from the Procedure Catalog screen.")

    screenshot(doc, "12-doctors.png",
               "The Doctors screen — the doctor directory and profiles.")

    doc.add_heading("4.6  Managing patients", level=2)
    para(doc,
         "Administrators and receptionists register patients and keep their records "
         "up to date. See Chapter 9 for the full patient workflow.")
    bullet(doc, "Register new patients with their personal and contact details.")
    bullet(doc, "Search and open any patient's profile.")
    bullet(doc, "Edit patient details.")
    bullet(doc, "Activate or deactivate a patient record — deactivated patients are excluded from most searches by default. Only an Administrator can do this.")

    doc.add_heading("4.7  Appointments, clinical records and treatment plans", level=2)
    para(doc,
         "As an administrator you can also manage appointments (Chapter 10), open "
         "and maintain patient records including diagnoses, prescriptions and "
         "follow-ups (Chapter 11), and manage treatment plans (Chapter 13). You are "
         "the only role that can delete (soft-delete) a patient record — the record "
         "is removed from active use while its history is kept — and, with the "
         "Chief Doctor, the only role that can view the audit trail of a patient "
         "record.")

    doc.add_heading("4.8  Billing", level=2)
    para(doc,
         "Administrators have the fullest access to billing: they can create and "
         "edit draft invoices, issue and cancel invoices, delete draft invoices, "
         "record and complete payments, allocate payments to invoices, generate and "
         "regenerate receipts, and run the full refund and credit note workflows. "
         "See Chapter 14 for the complete billing guide.")

    doc.add_heading("4.9  Administrative responsibilities", level=2)
    para(doc, "In day-to-day use, the administrator should normally:")
    bullet(doc, "Approve new staff accounts and assign the correct roles.")
    bullet(doc, "Create and maintain doctor profiles, specializations and schedules.")
    bullet(doc, "Keep the procedure catalogue up to date (prices, categories, statuses).")
    bullet(doc, "Activate or deactivate patient records when required.")
    bullet(doc, "Review the billing dashboard and outstanding amounts.")
    bullet(doc, "Handle actions restricted to administrators: deleting draft invoices or payments, deactivating users and patients, and finalising records.")

    doc.add_heading("4.10  What an administrator cannot do", level=2)
    bullet(doc, "Change your own role or deactivate your own account (the system blocks this deliberately).")
    bullet(doc, "Nothing else is blocked by design — but remember that every administrator action is recorded in the system's audit trail for accountability.")

    page_break(doc)

    # ── Section 5 ───────────────────────────────────────────────────
    doc.add_heading("5.  Doctor User Guide", level=1)

    para(doc,
         "This guide applies to General Doctors, Specialist Doctors, Consulting "
         "Doctors and Chief Doctors. Doctor-level accounts share the same clinical "
         "permissions. The Chief Doctor additionally has administrator-level access "
         "to management screens.")

    doc.add_heading("5.1  Signing in and the Dashboard", level=2)
    steps(doc, [
        "Open DensCare in your browser, enter your email and password, and click Sign In.",
        "On the Dashboard, use \u201cMy Treatment Plans\u201d to see the active treatment plans linked to you, and \u201cUpcoming Appointments\u201d to see today's appointments.",
        "Use the Quick Actions to jump straight to patients, appointments or invoice creation when needed.",
    ])

    doc.add_heading("5.2  Working with appointments", level=2)
    para(doc,
         "Open Appointments from the Clinical group. You can view the list, filter "
         "it, open an appointment's details, and reschedule or cancel appointments "
         "(with the other permitted staff). The system books appointments within "
         "clinic hours (10:00–13:00 and 17:00–21:00, Monday to Saturday) and checks "
         "for conflicts automatically.")

    doc.add_heading("5.3  Seeing patients", level=2)
    steps(doc, [
        "Open Patients and search for the patient by name, patient code or phone number.",
        "Open the patient's profile to see their contact details, emergency contact, alerts, allergies, clinical summary and treatment summary.",
        "To continue care, open the patient's clinical record (Patient Records) — you can also open an appointment to review the visit details.",
    ])

    doc.add_heading("5.4  Recording a consultation (patient record)", level=2)
    para(doc,
         "Each consultation is recorded as a patient record. Open Patient Records, "
         "create a new record for the patient, and fill in the clinical details: "
         "chief complaint, clinical notes, doctor remarks and treatment "
         "recommendation, plus the medical history fields (systemic diseases, "
         "surgeries, medications, habits, medical alerts, allergies and dental "
         "history).")
    para(doc,
         "A record moves through statuses: Draft → In Progress → Under Review → "
         "Completed → Finalized → Locked. Doctors (and administrators) update the "
         "status as the record matures. A finalized record is the final version for "
         "that visit.")

    doc.add_heading("5.5  Diagnoses", level=2)
    para(doc,
         "Inside a patient record, add one or more diagnoses. Each diagnosis has a "
         "name, a type — Provisional (initial) or Confirmed — and optional notes. "
         "Diagnoses can be edited or removed while the record is still open.")

    doc.add_heading("5.6  Prescriptions", level=2)
    steps(doc, [
        "Open the patient's record and go to the Prescriptions tab.",
        "Create a prescription and add up to 20 medicines. For each medicine enter the name, dosage (e.g. 500 mg), frequency (e.g. three times a day), duration (e.g. 5 days) and optional instructions (e.g. take after meals).",
        "Add any overall notes for the prescription.",
        "Save the prescription.",
        ("Print or download it for the patient", " — see Chapter 12 and Chapter 18."),
    ])
    para(doc,
         "Prescriptions can be edited while needed and are removed only by an "
         "administrator.")

    doc.add_heading("5.7  Follow-ups", level=2)
    para(doc,
         "Schedule follow-ups for the patient inside the record: record the "
         "follow-up date/time and notes. This helps the clinic remember the next "
         "visit.")

    doc.add_heading("5.8  Attachments", level=2)
    para(doc,
         "Upload supporting files — X-rays, scans, reports, consent forms — to a "
         "patient record from the Attachments tab. Files up to 10 MB in PDF, image "
         "and document formats are supported. You can preview or download them, "
         "change their category, or remove them.")

    doc.add_heading("5.9  Treatment plans", level=2)
    steps(doc, [
        "Open Treatment Plans and create a plan for the patient.",
        "Enter the clinical context: clinical notes, observations and dentist recommendations, and the plan's validity dates.",
        ("Add the procedures", " — pick procedures from the catalogue, specify the tooth (FDI number), quadrant and arch, and set the estimated cost."),
        "Work the plan through its statuses: Draft → Under Review → Proposed → Accepted → In Progress → Completed (it can also be Rejected, On Hold or Cancelled).",
        ("Obtain acceptance", " — record the patient's acceptance of the proposed plan."),
        "As treatment proceeds, update the plan status and each procedure's progress until the plan is Completed.",
    ])

    doc.add_heading("5.10  Billing", level=2)
    para(doc,
         "Doctors can create draft invoices for their patients, add line items "
         "(description, quantity, unit price, discount), issue and cancel invoices, "
         "record and complete payments, allocate payments to invoices, generate "
         "receipts, and create and process refunds and credit notes. The Billing "
         "Dashboard (Chapter 14) shows the financial picture. Only administrators "
         "can delete draft invoices or payments.")

    doc.add_heading("5.11  Your own profile", level=2)
    para(doc,
         "The system permits a doctor to view only their own doctor profile and "
         "schedule. The full doctor directory is available to administrators and "
         "receptionists — if you open a profile that is not yours, DensCare shows a "
         "permission message.")

    doc.add_heading("5.12  What doctors cannot do", level=2)
    bullet(doc, "Register new patients or edit patient details — this is done by the receptionist or administrator.")
    bullet(doc, "Manage user accounts or approve registrations (Chief Doctor can).")
    bullet(doc, "Manage the procedure catalogue or doctor profiles — these are administrator-level tasks (the Chief Doctor, as an administrator-level role, can).")
    bullet(doc, "Delete patient records, draft invoices or payments (administrators only).")
    bullet(doc, "View the patient record audit trail (administrator and Chief Doctor only).")
    bullet(doc, "View other doctors' profiles (only your own).")

    page_break(doc)

    # ── Section 6 ───────────────────────────────────────────────────
    doc.add_heading("6.  Receptionist User Guide", level=1)

    doc.add_heading("6.1  Signing in and the Dashboard", level=2)
    steps(doc, [
        "Open DensCare, enter your email and password, and click Sign In.",
        "Use the Dashboard to start the day: \u201cUpcoming Appointments\u201d shows today's schedule, and the Quick Actions take you to new-patient registration, appointment scheduling and invoice creation.",
    ])

    doc.add_heading("6.2  Patient registration", level=2)
    steps(doc, [
        "Open Patients and click New Patient (or use the New Patient quick action).",
        "Fill in the form: First name, Middle name (optional), Last name, Date of birth, Gender, Primary contact number, Emergency contact number (optional), Email address (optional), Address (optional) and Remarks (optional).",
        "Click Save Patient. DensCare assigns the patient a unique patient code automatically and adds them to the patient list.",
    ])

    doc.add_heading("6.3  Finding and updating patients", level=2)
    steps(doc, [
        "On the Patients screen, type a name, patient code or phone number into the search box.",
        "Use the status filter (All / Active / Inactive) if needed.",
        "Open the patient's profile to review their details.",
        "Click Edit on the profile to correct or update information and save.",
    ])

    doc.add_heading("6.4  Scheduling appointments", level=2)
    steps(doc, [
        "Open Appointments and click New Appointment (or use the Schedule Appointment quick action).",
        "Choose the patient (search by name or code), the dentist, and the appointment type: Consultation, Follow-Up, Emergency, Procedure, Review or Other.",
        "Pick the date (the clinic is closed on Sundays) and start time, and choose a duration of 15, 30, 45 or 60 minutes — the appointment must fall inside clinic hours (10:00–13:00 or 17:00–21:00).",
        "Enter the reason for the visit and any notes.",
        "Click Schedule Appointment. The system warns about conflicts with other appointments automatically.",
    ])

    doc.add_heading("6.5  Rescheduling and cancelling appointments", level=2)
    steps(doc, [
        "Open the appointment from the list and open its details.",
        ("Reschedule", " — use Edit, change the date, time or duration, and save. The patient cannot be changed after booking."),
        ("Cancel", " — click Cancel and confirm. A completed appointment cannot be cancelled."),
    ])

    doc.add_heading("6.6  Patient records", level=2)
    para(doc,
         "Receptionists can view patient records and add information such as "
         "records, diagnoses, prescriptions, follow-ups and attachments as needed, "
         "but only doctors (and administrators) can change a record's status or "
         "finalise it.")

    doc.add_heading("6.7  Billing at the counter", level=2)
    para(doc,
         "Receptionists handle most counter billing. You can create and issue "
         "invoices, record payments, allocate them to invoices, generate receipts, "
         "and create and process refunds and credit notes. Full procedures are in "
         "Chapter 14.")

    doc.add_heading("6.8  Your typical day", level=2)
    steps(doc, [
        "Sign in and check today's appointments on the Dashboard.",
        "Register any new patients who walk in.",
        "Schedule and reschedule appointments as patients call or arrive.",
        "Update patient information when patients report changes.",
        "Handle billing: create and issue invoices, record payments, give receipts.",
        "Complete the day's front-desk tasks, then sign out when you leave.",
    ])

    doc.add_heading("6.9  What receptionists cannot do", level=2)
    bullet(doc, "Activate or deactivate patient records (administrators only).")
    bullet(doc, "Change a patient record's status or finalise records (doctors and administrators only).")
    bullet(doc, "Delete patient records, draft invoices or payments (administrators only).")
    bullet(doc, "Manage user accounts, doctor profiles, schedules or the procedure catalogue (administrators only).")
    bullet(doc, "View the patient record audit trail (administrator and Chief Doctor only).")

    page_break(doc)

    # ── Section 7 ───────────────────────────────────────────────────
    doc.add_heading("7.  Dental Assistant User Guide", level=1)

    para(doc,
         "The Dental Assistant role is a billing-support role in the current "
         "version. Dental assistants help the front desk and the doctors with "
         "financial documents, but the system does not give them access to the "
         "clinical modules (patients, appointments, records, treatment plans) or to "
         "billing workflow actions.")

    doc.add_heading("7.1  Signing in", level=2)
    steps(doc, [
        "Open DensCare, enter your email and password, and click Sign In.",
        "You will land on the Dashboard, which shows the same quick actions and appointments information available to other staff.",
    ])

    doc.add_heading("7.2  What you can do in billing", level=2)
    para(doc, "Within the Financial group you can view everything and create the initial drafts of financial documents:")
    bullet(doc, "Billing Dashboard — view totals, recent invoices and payments.")
    bullet(doc, "Invoices — view, create and edit draft invoices.")
    bullet(doc, "Payments — view, create and edit payment records.")
    bullet(doc, "Receipts — view receipts and generate them for completed payments.")
    bullet(doc, "Refunds — create refund requests.")
    bullet(doc, "Credit notes — create draft credit notes.")

    doc.add_heading("7.3  What is restricted", level=2)
    para(doc,
         "A Dental Assistant can prepare documents but cannot complete the "
         "financial steps that require authorisation. The following actions are "
         "not available to this role:")
    bullet(doc, "Issuing or cancelling invoices, or deleting draft invoices.")
    bullet(doc, "Completing, failing or voiding payments, or allocating payments to invoices; deleting payments.")
    bullet(doc, "Regenerating receipts.")
    bullet(doc, "Approving, rejecting or completing refunds.")
    bullet(doc, "Issuing, applying or voiding credit notes.")
    para(doc,
         "If you open a screen your role cannot use — for example Patients or "
         "Appointments — DensCare shows a \u201cyou do not have permission\u201d message. "
         "Ask an administrator if you believe access should change.")

    page_break(doc)

    # ── Section 8 ───────────────────────────────────────────────────
    doc.add_heading("8.  DensCare Modules", level=1)

    para(doc,
         "This chapter explains every module in the delivered system. For each "
         "module: purpose, who uses it, what it does, main features, the typical "
         "workflow, important actions, related modules and notes.")

    # ── 8.1 Dashboard ──
    doc.add_heading("8.1  Dashboard", level=2)
    add_module_table(doc, [
        ("Purpose", "The home screen after sign-in; a starting point for the working day."),
        ("Who uses it", "Everyone with an account."),
        ("What it does", "Shows quick actions, the user's active treatment plans, and today's appointments."),
        ("Main features", "Overview statistics cards, four Quick Actions, My Treatment Plans, Upcoming Appointments, Recent Activity."),
        ("Typical workflow", "Sign in → check today's appointments → use a quick action to start a task."),
        ("Important actions", "New Patient, Schedule Appointment, Create Invoice, View Appointments."),
        ("Related modules", "Patients, Appointments, Invoices, Treatment Plans."),
        ("Notes", "The Overview numbers and Recent Activity list are sample/placeholder content in this version; live figures are on the Billing Dashboard."),
    ])

    # ── 8.2 Sign-in & account access ──
    doc.add_heading("8.2  Sign-in and account access", level=2)
    add_module_table(doc, [
        ("Purpose", "Secure entry to DensCare and account recovery."),
        ("Who uses it", "Everyone."),
        ("What it does", "Sign in with email and password, request a new account, recover a forgotten password."),
        ("Main features", "Sign in with \u201cRemember me\u201d, Create an account (self-registration), Forgot password (reset link by email), Reset password with a secure link."),
        ("Typical workflow", "Sign in with your email and password. New staff first request an account, wait for administrator approval, then sign in."),
        ("Important actions", "Sign In, Create an account, Forgot password."),
        ("Related modules", "Users & Roles (approval), All other modules."),
        ("Notes", "A new account is pending until an administrator approves it and assigns a role. Accounts that are deactivated cannot sign in."),
    ])

    # ── 8.3 Users & roles ──
    doc.add_heading("8.3  Users & roles", level=2)
    add_module_table(doc, [
        ("Purpose", "Manage staff accounts, roles and approvals."),
        ("Who uses it", "Administrators and Chief Doctors only."),
        ("What it does", "Lists users, shows each account's profile, changes roles, activates/deactivates accounts, approves pending registrations."),
        ("Main features", "User list with search/filters, user profile, role change, activate/deactivate, Pending Approvals screen."),
        ("Typical workflow", "Approve a new registration with a role → user signs in. Later, change roles or deactivate a leaver."),
        ("Important actions", "Approve, Change role, Activate, Deactivate."),
        ("Related modules", "Sign-in and account access, Doctors (doctor profiles link to user accounts)."),
        ("Notes", "You cannot change your own role or deactivate your own account."),
    ])

    # ── 8.4 Patients ──
    doc.add_heading("8.4  Patients", level=2)
    add_module_table(doc, [
        ("Purpose", "The central register of all clinic patients."),
        ("Who uses it", "Administrators and receptionists create and edit; doctors also view."),
        ("What it does", "Registers patients, keeps contact and medical alert details, provides each patient's profile."),
        ("Main features", "Patient list with search and status filter, registration form, patient profile, edit, activate/deactivate."),
        ("Typical workflow", "Register → search/open profile → edit details → (admin) activate/deactivate."),
        ("Important actions", "New Patient, Search, Edit, Activate, Deactivate."),
        ("Related modules", "Appointments, Patient Records, Treatment Plans, Billing (all reference the patient)."),
        ("Notes", "Patient codes are generated automatically. Deactivated patients are hidden from most searches by default."),
    ])

    # ── 8.5 Appointments ──
    doc.add_heading("8.5  Appointments", level=2)
    add_module_table(doc, [
        ("Purpose", "Schedule and manage patient visits."),
        ("Who uses it", "Administrators, receptionists and doctors."),
        ("What it does", "Books appointments within clinic hours, checks conflicts, tracks status, supports rescheduling and cancellation."),
        ("Main features", "New Appointment form, appointment list with status filter, Today's appointments, appointment details, edit, cancel."),
        ("Typical workflow", "Pick patient → pick dentist → choose date/time/duration → save. Update status as the visit progresses."),
        ("Important actions", "New Appointment, Edit (reschedule), Cancel, open details."),
        ("Related modules", "Patients, Doctors, Patient Records."),
        ("Notes", "Clinic hours are 10:00–13:00 and 17:00–21:00, Monday to Saturday (closed Sunday). Durations: 15/30/45/60 minutes."),
    ])

    # ── 8.6 Doctors ──
    doc.add_heading("8.6  Doctors", level=2)
    add_module_table(doc, [
        ("Purpose", "Doctor directory, specializations and weekly schedules."),
        ("Who uses it", "Administrators manage; receptionists view; doctors view their own profile."),
        ("What it does", "Stores doctor profiles linked to user accounts, specializations, availability, leave and weekly schedule templates."),
        ("Main features", "Doctor list, doctor profile, specialization assignment, schedule management, availability and leave toggles."),
        ("Typical workflow", "Admin creates profile for a user account → assigns specializations → sets the weekly schedule → doctor works; availability/leave toggles keep the list accurate."),
        ("Important actions", "Create doctor, Assign specializations, Set schedule, Toggle availability/leave, Deactivate."),
        ("Related modules", "Users, Appointments (dentist selection), Specializations."),
        ("Notes", "Doctors can only view their own profile and schedule; administrators and receptionists see the full directory."),
    ])

    # ── 8.7 Patient Records ──
    doc.add_heading("8.7  Patient Records", level=2)
    add_module_table(doc, [
        ("Purpose", "The clinical file for each patient — what happened at each visit."),
        ("Who uses it", "Doctors record and manage; receptionists and administrators also access."),
        ("What it does", "Stores consultation details, medical history, diagnoses, prescriptions, follow-ups and attachments, with a status lifecycle and audit trail."),
        ("Main features", "Record list, record detail with tabs (Clinical, Diagnoses, Prescriptions, Follow-ups, Attachments, Audit), status workflow, audit history."),
        ("Typical workflow", "Doctor opens/creates a record → records clinical details and diagnoses → adds a prescription → schedules a follow-up → finalises the record."),
        ("Important actions", "Create record, Edit clinical details, Change status, Finalize, add Diagnoses/Prescriptions/Follow-ups/Attachments, view Audit."),
        ("Related modules", "Patients, Appointments, Prescriptions, Attachments."),
        ("Notes", "Only administrators can delete a record (soft delete). Only administrators and Chief Doctors can view the audit trail."),
    ])

    # ── 8.8 Prescriptions ──
    doc.add_heading("8.8  Prescriptions", level=2)
    add_module_table(doc, [
        ("Purpose", "Write and manage medicines prescribed to patients."),
        ("Who uses it", "Doctors primarily; receptionists and administrators can also create them."),
        ("What it does", "Creates prescriptions with up to 20 medicines (name, dosage, frequency, duration, instructions) and prints/downloads them."),
        ("Main features", "Prescription list, medicine entry, notes, edit, printable document with Print / Download PDF."),
        ("Typical workflow", "Open the patient record → Prescriptions tab → add medicines → save → print for the patient."),
        ("Important actions", "Add medicine, Save, Open prescription (view/print/download), Edit, Delete (admin)."),
        ("Related modules", "Patient Records, Printing (Chapter 18)."),
        ("Notes", "Printing/downloading uses the browser's print dialog (choose \u201cSave as PDF\u201d to download)."),
    ])

    # ── 8.9 Treatment Plans ──
    doc.add_heading("8.9  Treatment Plans", level=2)
    add_module_table(doc, [
        ("Purpose", "Plan multi-step dental treatment and track its progress."),
        ("Who uses it", "Doctors create and manage; receptionists and administrators also access."),
        ("What it does", "Builds a plan from the procedure catalogue with tooth-level detail, manages its status lifecycle, records patient acceptance, and versions the plan as it changes."),
        ("Main features", "Plan list, plan detail with procedure items and cost summary, status actions, patient acceptance, version timeline."),
        ("Typical workflow", "Create plan → add procedures with teeth and costs → submit for review → propose to the patient → patient accepts → start treatment → complete."),
        ("Important actions", "Create plan, Add item, Submit for Review, Approve Review, Accept/Decline, Start Treatment, Hold, Resume, Complete, record patient acceptance."),
        ("Related modules", "Patients, Procedures (catalogue), Invoices (plan items can be billed)."),
        ("Notes", "Plans have versions — later changes create a new version instead of silently overwriting an accepted plan."),
    ])

    # ── 8.10 Procedure Catalogue ──
    doc.add_heading("8.10  Procedure Catalogue", level=2)
    add_module_table(doc, [
        ("Purpose", "The master list of dental procedures used in treatment plans and billing."),
        ("Who uses it", "Administrators manage; doctors, receptionists and administrators use it when planning treatment."),
        ("What it does", "Stores each procedure's name, category, price and status."),
        ("Main features", "Procedure list with search/filter, procedure creation/edit, activate/deactivate."),
        ("Typical workflow", "Admin adds/updates procedures and prices → doctors pick them when building treatment plans."),
        ("Important actions", "Add procedure, Edit, Activate/Deactivate."),
        ("Related modules", "Treatment Plans, Invoices."),
        ("Notes", "Only administrator-level accounts (Administrator and Chief Doctor) can create or change procedures."),
    ])

    screenshot(doc, "11-procedures.png",
               "The Procedure Catalogue — the master list of dental procedures and prices.")

    # ── 8.11 Billing ──
    doc.add_heading("8.11  Billing (Invoices, Payments, Receipts, Refunds, Credit Notes)", level=2)
    add_module_table(doc, [
        ("Purpose", "The clinic's financial engine: invoicing, payments, receipts, refunds and credit notes."),
        ("Who uses it", "Receptionists, administrators, doctors and dental assistants (with the limits in Chapter 22)."),
        ("What it does", "Creates and issues invoices, records and allocates payments, generates receipts, processes refunds and applies credit notes, and summarises the financial position."),
        ("Main features", "Billing Dashboard, Invoice list/detail, Payment list/detail, Receipt detail, Refund timeline, Credit note detail, printable invoices."),
        ("Typical workflow", "Create invoice → issue it → record payment → allocate to the invoice → generate receipt → (if needed) refund or credit note."),
        ("Important actions", "Create/Issue/Cancel invoice, Record/Complete payment, Allocate payment, Generate receipt, Create/Approve/Complete refund, Create/Issue/Apply credit note, Print/Download invoice."),
        ("Related modules", "Patients, Treatment Plans, Appointments."),
        ("Notes", "Currencies supported: USD, EUR, GBP, INR. The interface presents amounts in Indian Rupees (₹). Refunds and credit notes are created from the payment/invoice detail screens."),
    ])

    doc.add_heading("8.12  Modules not yet available", level=2)
    para(doc,
         "Inventory, Laboratory, Reports and Settings appear in the navigation as "
         "disabled placeholders and are not functional in this version. See "
         "Chapter 26 for the full list of planned-but-unavailable features.")

    page_break(doc)

    # ── Section 9 ───────────────────────────────────────────────────
    doc.add_heading("9.  Patient Management — Step-by-Step", level=1)

    doc.add_heading("9.1  Register a new patient", level=2)
    steps(doc, [
        "Open Patients from the Clinical group in the left menu.",
        "Click New Patient (top-right of the list).",
        "Complete the form. Required fields are marked: First name, Last name, Date of birth, Gender and Primary contact number.",
        ("Useful entry rules", " — names may contain letters, spaces, hyphens and apostrophes (2–100 characters); phone numbers must be 10–15 digits with an optional leading \u201c+\u201d; the date of birth cannot be in the future."),
        "Click Save Patient.",
        "The patient appears in the list with a system-generated patient code. You can open their profile immediately.",
    ])
    callout(doc, "Tip",
            "Enter a phone number as your primary search key — the patient search "
            "matches patient code, name and phone number.")

    screenshot(doc, "03-patients.png",
               "The Patients list — search, filter and open patient profiles.")

    doc.add_heading("9.2  Search for a patient", level=2)
    steps(doc, [
        "Open Patients.",
        "Type a name, patient code or phone number in the search box — the list updates as you type.",
        "Browse the results and click a patient to open their profile.",
    ])

    doc.add_heading("9.3  Filter the patient list", level=2)
    steps(doc, [
        "Open Patients.",
        "Use the status filter to show All patients, only Active patients, or only Inactive patients.",
        "Use the page controls at the bottom of the list to move through more than one page of results (the list shows up to 20 patients per page).",
    ])

    doc.add_heading("9.4  Open a patient profile", level=2)
    para(doc,
         "A patient profile brings everything about the patient together: personal "
         "information, contact details, emergency contact, status, and summary "
         "cards for upcoming appointments, treatment plans, clinical details, "
         "allergies and alerts.")

    screenshot(doc, "04-patient-profile.png",
               "A patient profile — personal details, contacts and clinical summary.")

    doc.add_heading("9.5  Edit a patient", level=2)
    steps(doc, [
        "Open the patient's profile.",
        "Click Edit.",
        "Update the required fields and save. Duplicate checks prevent exact duplicates when you change names or phone numbers.",
    ])

    doc.add_heading("9.6  Activate / deactivate a patient (Administrator only)", level=2)
    steps(doc, [
        "Open the patient's profile.",
        ("Deactivate", " — use the deactivate action. Deactivated patients are excluded from most searches by default but their history is kept."),
        ("Activate", " — use the activate action to bring the patient back into the active register."),
    ])

    doc.add_heading("9.7  View patient history", level=2)
    para(doc,
         "From the patient's profile you can jump to their appointments, clinical "
         "records, treatment plans and billing documents. Each of those is stored "
         "in its own module but is always linked back to the patient.")

    page_break(doc)

    # ── Section 10 ──────────────────────────────────────────────────
    doc.add_heading("10.  Appointment Management — Step-by-Step", level=1)

    doc.add_heading("10.1  Create an appointment", level=2)
    steps(doc, [
        "Open Appointments and click New Appointment (or use the Schedule Appointment quick action on the Dashboard).",
        "Search and select the patient.",
        "Select the dentist.",
        "Choose the appointment type: Consultation, Follow-Up, Emergency, Procedure, Review or Other.",
        "Pick the date (Monday to Saturday — the clinic is closed on Sundays) and the start time.",
        "Choose the duration: 15, 30, 45 or 60 minutes. The appointment must fit inside a clinic session (10:00–13:00 or 17:00–21:00).",
        "Enter the reason for the visit (3–500 characters) and any notes.",
        "Click Schedule Appointment. DensCare checks for double-booking automatically and warns you if the slot conflicts.",
    ])

    screenshot(doc, "05-appointments.png",
               "The Appointments screen — schedule, reschedule and cancel visits.")

    doc.add_heading("10.2  View today's appointments", level=2)
    steps(doc, [
        "Open Appointments and use the date filter to show Today, or check the Upcoming Appointments panel on the Dashboard.",
    ])

    doc.add_heading("10.3  Search and filter appointments", level=2)
    para(doc,
         "The appointment list has a status filter (All plus each status) so you "
         "can, for example, show only today's scheduled or completed appointments. "
         "Search and pagination work like every other list in DensCare (Chapter 16).")

    doc.add_heading("10.4  Reschedule an appointment", level=2)
    steps(doc, [
        "Open the appointment from the list.",
        "Click Edit.",
        "Change the date, time or duration and save. The patient cannot be changed after booking.",
    ])

    doc.add_heading("10.5  Cancel an appointment", level=2)
    steps(doc, [
        "Open the appointment from the list.",
        "Click Cancel and confirm. A Completed appointment cannot be cancelled.",
    ])

    doc.add_heading("10.6  Appointment details and status", level=2)
    para(doc,
         "Opening an appointment shows the patient, dentist, schedule, type, "
         "reason and notes, together with the current status badge. The status "
         "reflects where the visit is in its lifecycle:")
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["Scheduled", "The appointment has been booked."],
            ["Confirmed", "The appointment has been confirmed."],
            ["Checked In", "The patient has arrived at the clinic."],
            ["In Treatment", "The patient is currently with the doctor."],
            ["Completed", "The visit is finished."],
            ["Cancelled", "The appointment was cancelled."],
            ["No Show", "The patient did not arrive."],
        ],
        col_widths=[Cm(3.5), Cm(12.5)],
    )

    page_break(doc)

    # ── Section 11 ──────────────────────────────────────────────────
    doc.add_heading("11.  Patient Records — Step-by-Step", level=1)

    doc.add_heading("11.1  Clinical records", level=2)
    para(doc, "A patient record is the clinical file for one visit or episode of care.")
    steps(doc, [
        "Open Patient Records and create a record for the patient (or open an existing one).",
        ("Fill in the clinical section", " — Chief Complaint, Clinical Notes, Doctor Remarks, Treatment Recommendation."),
        ("Fill in the medical history section", " — Systemic Diseases, Surgeries, Medications, Habits, Medical Alerts, Allergies, Dental History."),
        "Save the record.",
        "Update the record status as care progresses: Draft → In Progress → Under Review → Completed → Finalized (→ Locked). Doctors and administrators can change status; finalisation marks the record as final.",
    ])

    screenshot(doc, "06-patient-records.png",
               "The Patient Records list — one record per consultation.")
    screenshot(doc, "07-record-detail.png",
               "A patient record — the Clinical tab with complaints, notes and medical history.")

    doc.add_heading("11.2  Diagnoses", level=2)
    steps(doc, [
        "Open the patient's record and go to the Diagnoses tab.",
        "Add a diagnosis: enter the name, choose the type (Provisional for an initial impression, Confirmed for a firm diagnosis) and add optional notes.",
        "Save. You can add several diagnoses to one record and edit or remove them later.",
    ])

    doc.add_heading("11.3  Prescriptions", level=2)
    steps(doc, [
        "Open the patient's record and go to the Prescriptions tab.",
        "Create a prescription, then add the medicines (up to 20 per prescription). For each medicine enter: Medicine name, Dosage, Frequency, Duration, Instructions (optional).",
        "Add overall prescription notes if needed and save.",
        "Open the prescription to view, print or download it (Chapter 18).",
    ])

    screenshot(doc, "08a-record-prescriptions-tab.png",
               "The Prescriptions tab inside a patient record.")

    doc.add_heading("11.4  Follow-ups", level=2)
    steps(doc, [
        "Open the patient's record and go to the Follow-ups tab.",
        "Add a follow-up with the follow-up details and notes, and save. Follow-ups remind the clinic to bring the patient back.",
    ])

    doc.add_heading("11.5  Attachments", level=2)
    steps(doc, [
        "Open the patient's record and go to the Attachments tab.",
        "Upload files (up to 10 MB; PDF, JPG, PNG, GIF, WEBP, TIFF, BMP, DOC, DOCX, TXT) and choose the attachment type: Image, PDF, Report, Scan or Document.",
        "Preview or download the file, change its category, or remove it.",
    ])

    doc.add_heading("11.6  Audit / history", level=2)
    para(doc,
         "Every change to a patient record and its diagnoses, prescriptions, "
         "follow-ups and attachments is recorded in the audit trail — who did what "
         "and when. The Audit tab is visible to Administrators and Chief Doctors "
         "only, so records stay accountable without exposing internal history to "
         "everyone.")

    page_break(doc)

    # ── Section 12 ──────────────────────────────────────────────────
    doc.add_heading("12.  Prescriptions — Step-by-Step", level=1)

    doc.add_heading("12.1  Create a prescription", level=2)
    steps(doc, [
        "Open the patient's record and go to the Prescriptions tab.",
        "Click to create a new prescription.",
        "Add each medicine with its Name, Dosage, Frequency, Duration and optional Instructions.",
        "Add any notes for the prescription as a whole.",
        "Save the prescription.",
    ])

    doc.add_heading("12.2  Add or edit a medicine", level=2)
    para(doc,
         "Each medicine entry can be added, edited or removed while the "
         "prescription is being prepared. A prescription can contain up to 20 "
         "medicines.")

    doc.add_heading("12.3  View a prescription", level=2)
    steps(doc, [
        "Open the patient's record → Prescriptions tab.",
        "Click the prescription to open it in a printable preview with all medicines and instructions.",
    ])

    doc.add_heading("12.4  Print or download a prescription", level=2)
    steps(doc, [
        "Open the prescription (step 12.3).",
        ("Print", " — click Print. Your browser's print dialog opens; choose your printer and print."),
        ("Download PDF", " — click Download PDF. Your browser's print dialog opens; choose \u201cSave as PDF\u201d as the destination and save."),
    ])
    callout(doc, "Note",
            "DensCare does not silently generate a PDF file — it opens the "
            "browser's print dialog, and \u201cSave as PDF\u201d is chosen there. This is "
            "the intended behaviour of the current version.")

    screenshot(doc, "08b-prescription-printable.png",
               "A printable prescription — medicines, dosage and instructions.")

    doc.add_heading("12.5  Edit or delete a prescription", level=2)
    para(doc,
         "Prescriptions can be edited (e.g. adjust a dosage) while the record is "
         "open. Removing a prescription entirely is restricted to administrators.")

    page_break(doc)

    # ── Section 13 ──────────────────────────────────────────────────
    doc.add_heading("13.  Treatment Plans — Step-by-Step", level=1)

    doc.add_heading("13.1  Create a treatment plan", level=2)
    steps(doc, [
        "Open Treatment Plans and click to create a new plan.",
        "Select the patient and the doctor.",
        "Enter the clinical context: Clinical Notes, Observations, and Dentist Recommendations (each up to 5,000 characters).",
        "Optionally set the validity dates (Valid From / Valid To) and a plan code.",
        "Create the plan. It starts in Draft status.",
    ])

    screenshot(doc, "09-treatment-plans.png",
               "The Treatment Plans list.")

    doc.add_heading("13.2  Add procedures to the plan", level=2)
    steps(doc, [
        "Open the plan and add an item.",
        "Choose a procedure from the catalogue.",
        "Enter the tooth-level detail: tooth number (FDI notation, e.g. 11–48 permanent or 51–85 primary), quadrant (UR, UL, LL, LR) and arch (upper/lower).",
        "Set the estimated cost and any item notes.",
        "Add as many items as the treatment requires; the plan shows the total.",
    ])

    doc.add_heading("13.3  Manage the plan's status", level=2)
    para(doc,
         "A plan moves through a defined lifecycle. The available actions appear "
         "on the plan's detail screen according to its current status:")
    add_table(
        doc,
        ["Status", "Typical meaning"],
        [
            ["Draft", "The plan is being prepared; not yet shared."],
            ["Under Review", "The plan has been submitted for review by the doctor/team."],
            ["Proposed", "The plan has been approved internally and is proposed to the patient."],
            ["Accepted", "The patient has accepted the plan."],
            ["In Progress", "Treatment has started and is being carried out."],
            ["On Hold", "Treatment is temporarily paused."],
            ["Completed", "All planned treatment is finished."],
            ["Rejected", "The plan was rejected during review."],
            ["Cancelled", "The plan was cancelled and will not proceed."],
        ],
        col_widths=[Cm(3.5), Cm(12.5)],
    )
    para(doc,
         "Typical actions: Submit for Review → Approve Review → Accept Plan (record "
         "the patient's acceptance) → Start Treatment → Complete Treatment. Plans "
         "can also be declined, cancelled, put on hold or resumed at the right "
         "points. Because accepted plans are versioned, any later change creates a "
         "new version rather than silently rewriting the agreed plan.")

    screenshot(doc, "10-treatment-plan-detail.png",
               "A treatment plan detail — procedures, costs and status.")

    doc.add_heading("13.4  Relationship with the patient and billing", level=2)
    bullet(doc, "Patient — a plan always belongs to one patient and appears in the patient's profile and in the doctor's \u201cMy Treatment Plans\u201d on the Dashboard.")
    bullet(doc, "Billing — plan items can be carried into an invoice, so agreed procedures become the line items you bill.")

    page_break(doc)

    # ── Section 14 ──────────────────────────────────────────────────
    doc.add_heading("14.  Billing — Step-by-Step", level=1)

    doc.add_heading("14.1  Billing Dashboard", level=2)
    para(doc,
         "The Billing Dashboard is the clinic's financial overview. It shows "
         "system-wide totals — total invoiced, collected, outstanding, refunded and "
         "credit notes — together with the five most recent invoices and the five "
         "most recent payments. You can also filter the view to a single patient to "
         "see that patient's financial summary.")

    screenshot(doc, "13-billing-dashboard.png",
               "The Billing Dashboard — the clinic's financial overview.")

    doc.add_heading("14.2  Invoices", level=2)

    doc.add_heading("14.2.1  Create an invoice", level=3)
    steps(doc, [
        "Open Invoices and click to create a new invoice (or use the Create Invoice quick action on the Dashboard).",
        "Select the patient. Optionally link the invoice to a treatment plan, an appointment and a doctor.",
        "Choose the currency (the interface presents amounts in Indian Rupees, ₹).",
        "Set the invoice date and the due date (due date cannot be before the invoice date).",
        "Add at least one line item: Description, Quantity (at least 1), Unit price, and optionally a discount (Percentage or Fixed amount). The discount cannot exceed the line subtotal.",
        "Add notes if required and save. The invoice is created as a Draft with a temporary number.",
    ])

    screenshot(doc, "14-invoices.png",
               "The Invoices list.")

    doc.add_heading("14.2.2  Issue an invoice", level=3)
    steps(doc, [
        "Open the draft invoice.",
        "Click Issue. DensCare assigns a permanent invoice number and the invoice becomes a financial document — from this point its amounts are fixed.",
    ])

    doc.add_heading("14.2.3  Edit a draft invoice", level=3)
    para(doc,
         "While an invoice is still a Draft you can edit its notes and due date. "
         "Once issued, an invoice is immutable — corrections are handled by "
         "cancellation and a new invoice, or by a credit note.")

    doc.add_heading("14.2.4  Invoice status", level=3)
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["Draft", "Being prepared; not yet a financial document."],
            ["Issued", "Sent to the patient with a permanent number; awaiting payment."],
            ["Partially Paid", "Some of the total has been paid."],
            ["Paid", "Fully settled."],
            ["Overdue", "Past the due date with an outstanding balance."],
            ["Cancelled", "Cancelled with a recorded reason."],
            ["Void", "Voided and without financial effect."],
        ],
        col_widths=[Cm(3.5), Cm(12.5)],
    )

    doc.add_heading("14.2.5  Cancel an invoice", level=3)
    steps(doc, [
        "Open the invoice and click Cancel.",
        "Enter the reason for cancellation (required).",
        "Confirm. The invoice moves to Cancelled status.",
    ])

    doc.add_heading("14.2.6  Print or download an invoice", level=3)
    steps(doc, [
        "Open the invoice.",
        "Click Print or Download PDF — the browser's print dialog opens.",
        "Choose your printer, or choose \u201cSave as PDF\u201d as the destination to download.",
    ])

    screenshot(doc, "15-invoice-detail.png",
               "An invoice detail — line items, totals and payment status.")

    doc.add_heading("14.2.7  Delete a draft invoice (Administrator only)", level=3)
    para(doc,
         "Only administrators can permanently delete an invoice, and only while it "
         "is still a Draft. Issued invoices are never deleted.")

    doc.add_heading("14.3  Payments", level=2)

    doc.add_heading("14.3.1  Record a payment", level=3)
    steps(doc, [
        "Open Payments and click to record a new payment.",
        "Select the patient, enter the amount (greater than 0) and choose the payment method: Cash, Card, UPI, Bank Transfer, Cheque, Insurance or Wallet.",
        "Set the payment date, and optionally a reference number and notes.",
        "Save. The payment is created in Pending status.",
    ])

    screenshot(doc, "16-payments.png",
               "The Payments list.")

    doc.add_heading("14.3.2  Complete and allocate a payment", level=3)
    steps(doc, [
        "Open the payment and click Complete — the payment becomes Completed.",
        ("Allocate the payment to an invoice", " — choose the invoice and the amount to allocate. A payment can be split across several invoices."),
        "Once an invoice's outstanding balance is fully covered it becomes Paid; partial allocation shows Partially Paid.",
    ])

    screenshot(doc, "17-payment-detail.png",
               "A payment detail — status, allocation and totals.")

    doc.add_heading("14.3.3  Payment status", level=3)
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["Pending", "Recorded but not yet completed."],
            ["Completed", "The payment is confirmed and can be allocated to invoices."],
            ["Failed", "The payment did not go through (reason recorded)."],
            ["Refunded", "The payment has been fully refunded."],
            ["Reversed", "The payment was reversed."],
            ["Void", "The payment was voided (reason recorded)."],
        ],
        col_widths=[Cm(3.5), Cm(12.5)],
    )

    doc.add_heading("14.3.4  Edit, fail, void or delete a payment", level=3)
    steps(doc, [
        ("Edit", " — a Pending payment's reference number and notes can be updated."),
        ("Fail / Void", " — mark the payment as failed or void with an optional reason."),
        ("Delete", " — only administrators can permanently delete a payment, and only while it is Pending."),
    ])

    doc.add_heading("14.4  Receipts", level=2)
    steps(doc, [
        ("Generate a receipt", " — for a Completed payment, generate its receipt. DensCare assigns a permanent receipt number."),
        ("View a receipt", " — open it to see the patient, payment, amount and dates."),
        ("Regenerate", " — reproduce the same receipt (for a lost or reprinted copy) without creating a new financial record."),
        ("Print / download", " — use the browser print dialog as described in Chapter 18."),
    ])

    screenshot(doc, "18-receipt.png",
               "A receipt generated for a completed payment.")

    doc.add_heading("14.5  Refunds", level=2)
    para(doc,
         "Refunds return money to a patient. A refund is created against a "
         "completed payment and must be approved before it can be completed. Only "
         "the amount not yet refunded can be refunded.")
    steps(doc, [
        ("Create the refund", " — from the payment's detail screen, open Create Refund, enter the amount (up to the refundable balance) and the reason. The refund starts in Pending status."),
        ("Approve the refund", " — an authorised staff member approves it."),
        ("Complete the refund", " — the refund allocation is created; if the whole payment is refunded, the payment shows the Refunded status."),
        ("Reject", " — a pending refund can be rejected with a reason if it should not proceed."),
    ])
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["Pending", "Requested, awaiting approval."],
            ["Approved", "Approved; ready to be completed."],
            ["Completed", "The refund has been processed (terminal)."],
            ["Rejected", "The request was declined with a reason (terminal)."],
        ],
        col_widths=[Cm(3.5), Cm(12.5)],
    )

    doc.add_heading("14.6  Credit Notes", level=2)
    para(doc,
         "A credit note is a document that gives a patient credit against an "
         "invoice — typically used for corrections, partial cancellations or "
         "goodwill adjustments.")
    steps(doc, [
        ("Create the credit note", " — from the invoice's detail screen, open Create Credit Note, select the invoice and patient, enter the amount (up to the invoice total), a reason, and optionally an expiry date. It starts in Draft status."),
        ("Issue the credit note", " — a permanent credit note number is assigned."),
        ("Apply the credit note", " — apply the credit to the invoice; its remaining balance becomes zero."),
        ("Void", " — an unapplied credit note can be voided with a reason."),
    ])
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["Draft", "Being prepared."],
            ["Issued", "Issued with a permanent number; can be applied or voided."],
            ["Applied", "The credit has been used."],
            ["Void", "Voided with a reason."],
            ["Expired", "Reached its expiry date without being applied."],
        ],
        col_widths=[Cm(3.5), Cm(12.5)],
    )

    page_break(doc)

    # ── Section 15 ──────────────────────────────────────────────────
    doc.add_heading("15.  File Attachments — Step-by-Step", level=1)

    doc.add_heading("15.1  Upload a file", level=2)
    steps(doc, [
        "Open the patient's record and go to the Attachments tab.",
        "Choose Upload and select the file from your computer (or drag it into the upload area).",
        "Choose the attachment type: Image, PDF, Report, Scan or Document.",
        "Confirm the upload. The file is added to the patient record.",
    ])

    doc.add_heading("15.2  Supported file types and size", level=2)
    para(doc, "DensCare accepts these formats, up to 10 MB per file:")
    para(doc, "PDF, JPG, JPEG, PNG, GIF, WEBP, TIFF, BMP, DOC, DOCX, TXT", bold=True)

    doc.add_heading("15.3  View / preview", level=2)
    para(doc,
         "PDF and image files can be previewed directly in the browser. For other "
         "file types, use Download instead — the preview may not be available.")

    doc.add_heading("15.4  Download", level=2)
    steps(doc, [
        "Open the patient record → Attachments tab.",
        "Click the download icon on the file. The file is saved to your computer with its original name.",
    ])

    doc.add_heading("15.5  Delete", level=2)
    steps(doc, [
        "Open the patient record → Attachments tab.",
        "Click the delete action on the file and confirm.",
    ])

    doc.add_heading("15.6  Who can manage attachments", level=2)
    para(doc,
         "Staff who can work with patient records (administrators, receptionists "
         "and doctors) can upload, preview, download, recategorise and remove "
         "attachments. Every upload and download is recorded in the record's audit "
         "trail.")

    page_break(doc)

    # ── Section 16 ──────────────────────────────────────────────────
    doc.add_heading("16.  Search, Filters and Tables", level=1)

    doc.add_heading("16.1  Search", level=2)
    para(doc,
         "Most list screens (Patients, Appointments, Patient Records, Treatment "
         "Plans, Invoices, Payments, Users, Doctors, Procedures) have a search box. "
         "Type a few characters and the list narrows immediately. What is matched "
         "depends on the module — for example, patient search matches name, patient "
         "code and phone number, while invoice search matches the invoice number "
         "and patient name.")

    doc.add_heading("16.2  Filters", level=2)
    para(doc,
         "Lists offer drop-down filters for the things you search by most: status "
         "(e.g. Active/Inactive patients, appointment status, invoice/payment "
         "status), type or method where relevant, and in the billing lists date "
         "ranges. Filters combine with search, so you can, for example, show all "
         "Overdue invoices for one patient within a date range.")

    doc.add_heading("16.3  Sorting", level=2)
    para(doc,
         "Where sorting is available (for example the invoice and payment lists) "
         "you can sort by fields such as document number, status, date or amount, "
         "in ascending or descending order.")

    doc.add_heading("16.4  Pagination", level=2)
    para(doc,
         "Lists show a limited number of rows per page (20 by default). Use the "
         "page controls at the bottom to move between pages. In most lists you can "
         "choose how many rows to show per page (up to 100).")

    doc.add_heading("16.5  Column controls", level=2)
    para(doc,
         "Some list screens — Patients, Appointments, Doctors, Users and the "
         "Procedure Catalogue — let you show or hide columns using the Columns "
         "control in the toolbar, so you can focus on the fields you use most. "
         "Other lists do not offer this control.")

    doc.add_heading("16.6  Detail pages", level=2)
    para(doc,
         "Clicking a row opens that item's detail page, which shows the full "
         "record and the actions available for its current state.")

    page_break(doc)

    # ── Section 17 ──────────────────────────────────────────────────
    doc.add_heading("17.  The Dashboard in Detail", level=1)

    doc.add_heading("17.1  Overview statistics", level=2)
    para(doc,
         "Four cards — Total Patients, Today's Appointments, Revenue Today and "
         "Pending Treatments. In the current version these cards show sample "
         "figures; the live numbers are available on the Billing Dashboard.")

    doc.add_heading("17.2  Quick Actions", level=2)
    para(doc, "Four buttons that jump straight to a task:")
    add_table(
        doc,
        ["Quick action", "What it does", "Who can use it"],
        [
            ["New Patient", "Opens the Patients screen with the registration form ready to fill in.", "Administrators and receptionists."],
            ["Schedule Appointment", "Opens the Appointments screen with the booking form ready.", "Administrators, receptionists and doctors."],
            ["Create Invoice", "Opens the Invoices screen with the new-invoice form ready.", "All billing roles (see Chapter 22)."],
            ["View Appointments", "Opens the Appointments list.", "All clinical roles."],
        ],
        col_widths=[Cm(4.4), Cm(7.0), Cm(4.6)],
    )

    doc.add_heading("17.3  My Treatment Plans", level=2)
    para(doc,
         "The live list of treatment plans linked to the signed-in doctor. It "
         "shows each plan's patient and status so a doctor can see what is in "
         "progress at a glance.")

    doc.add_heading("17.4  Upcoming Appointments", level=2)
    para(doc,
         "A live view of today's appointments with patient, dentist and status.")

    doc.add_heading("17.5  Recent Activity", level=2)
    para(doc,
         "A sample activity feed in the current version; the authoritative, "
         "up-to-date records are in each module's own list.")

    page_break(doc)

    # ── Section 18 ──────────────────────────────────────────────────
    doc.add_heading("18.  Printing and Documents", level=1)

    para(doc,
         "DensCare produces professional, printable documents for invoices and "
         "prescriptions. Printing is handled through your browser's own print "
         "dialog.")

    doc.add_heading("18.1  Invoice printing", level=2)
    steps(doc, [
        "Open the invoice.",
        "Click Print (or Download PDF). DensCare opens a preview of the invoice and then your browser's print dialog.",
        ("To print", " — choose your printer and click Print."),
        ("To save as PDF", " — choose \u201cSave as PDF\u201d as the destination in the print dialog and save."),
    ])

    doc.add_heading("18.2  Prescription printing", level=2)
    steps(doc, [
        "Open the prescription from the patient record (Prescriptions tab).",
        "Click Print or Download PDF.",
        "Use the browser print dialog to print or save as PDF, as above.",
    ])

    doc.add_heading("18.3  Receipts", level=2)
    para(doc,
         "Receipts are viewed in the same way; use your browser's print dialog to "
         "print a receipt for the patient.")

    callout(doc, "Note — how PDF download works",
            "\u201cDownload PDF\u201d does not silently create a file: it opens the browser's "
            "print dialog, where you choose \u201cSave as PDF\u201d as the destination. This "
            "is the intended behaviour of the current version.")

    page_break(doc)

    # ── Section 19 ──────────────────────────────────────────────────
    doc.add_heading("19.  Mobile Usage", level=1)

    para(doc,
         "DensCare works on phones and tablets as well as computers. The screens "
         "adapt to the screen size automatically.")

    doc.add_heading("19.1  Navigation on a phone", level=2)
    bullet(doc, "The left sidebar is hidden on small screens. Tap the menu button in the header to open it as a slide-in drawer.")
    bullet(doc, "On the Dashboard, a bottom navigation bar gives quick access to the main areas.")
    bullet(doc, "On list screens, the header shows a compact layout with the menu button, the page title and the main action (such as New Patient).")

    doc.add_heading("19.2  Lists become cards", level=2)
    para(doc,
         "Tables that fit on a computer screen are shown as cards on a phone — "
         "each patient, appointment, invoice or payment becomes its own card with "
         "the key information and a status badge.")

    screenshot(doc, "21-mobile-patients.png",
               "The Patients screen on a phone — list rows become cards.",
               width_cm=8.5)

    doc.add_heading("19.3  Forms, drawers and dialogs", level=2)
    para(doc,
         "Forms and creation screens open in drawers or dialogs on all devices and "
         "are fully usable on a phone. Fields stack vertically, so allow a little "
         "scrolling for longer forms.")

    doc.add_heading("19.4  Printing from a phone", level=2)
    para(doc,
         "Print and Download PDF work the same way on a phone: the browser's print "
         "dialog opens, where you can print to a printer or save as PDF.")

    doc.add_heading("19.5  Differences to keep in mind", level=2)
    bullet(doc, "More scrolling and tapping — cards and drawers replace the desktop tables and side-by-side layouts.")
    bullet(doc, "The full sidebar is always reachable through the menu button, so nothing is missing on mobile.")

    page_break(doc)

    # ── Section 20 ──────────────────────────────────────────────────
    doc.add_heading("20.  Common Daily Workflows", level=1)

    doc.add_heading("Scenario 1 — A new patient's first visit", level=2)
    steps(doc, [
        ("Receptionist:", " register the patient (Chapter 9.1)."),
        ("Receptionist:", " schedule an appointment with a doctor (Chapter 10.1)."),
        ("Doctor:", " open the appointment and review the patient."),
        ("Doctor:", " create the patient record and record the consultation and diagnosis (Chapter 11)."),
        ("Doctor:", " create a treatment plan if treatment is planned, and obtain patient acceptance (Chapter 13)."),
        ("Doctor:", " write the prescription (Chapter 12)."),
        ("Receptionist:", " create and issue the invoice (Chapter 14.2)."),
        ("Receptionist:", " record the payment and allocate it to the invoice (Chapter 14.3)."),
        ("Receptionist:", " generate the receipt and hand it to the patient (Chapter 14.4)."),
    ])

    doc.add_heading("Scenario 2 — An existing patient returns for a follow-up", level=2)
    steps(doc, [
        ("Receptionist:", " search for the patient and open their profile."),
        ("Receptionist:", " book a Follow-Up appointment."),
        ("Doctor:", " open the patient's previous record to review history."),
        ("Doctor:", " record the follow-up visit and schedule the next follow-up."),
    ])

    doc.add_heading("Scenario 3 — The patient needs a prescription", level=2)
    steps(doc, [
        ("Doctor:", " open the patient's record → Prescriptions tab."),
        ("Doctor:", " create the prescription, add the medicines with dosage, frequency, duration and instructions."),
        ("Doctor:", " save and print the prescription (or save it as PDF) for the patient (Chapter 12)."),
    ])

    doc.add_heading("Scenario 4 — The patient makes a payment", level=2)
    steps(doc, [
        ("Receptionist:", " open Payments and record the payment with the method and amount."),
        ("Receptionist:", " complete the payment."),
        ("Receptionist:", " allocate the payment to the patient's invoice."),
        ("Receptionist:", " generate and print the receipt."),
    ])

    doc.add_heading("Scenario 5 — A refund", level=2)
    steps(doc, [
        ("Receptionist:", " open the relevant payment and create a refund for the refundable amount with a reason (Chapter 14.5)."),
        ("Authorised staff:", " approve the refund."),
        ("Authorised staff:", " complete the refund."),
    ])

    doc.add_heading("Scenario 6 — The patient needs a copy of an invoice", level=2)
    steps(doc, [
        ("Receptionist:", " open Invoices and search for the patient or invoice number."),
        ("Receptionist:", " open the invoice and click Print or Download PDF (Chapter 18)."),
    ])

    page_break(doc)

    # ── Section 21 ──────────────────────────────────────────────────
    doc.add_heading("21.  Frequently Asked Questions (FAQ)", level=1)

    faqs = [
        ("How do I register a new patient?",
         "Open Patients → New Patient, complete the form and click Save Patient (Chapter 9.1)."),
        ("How do I search for a patient?",
         "Open Patients and type the name, patient code or phone number into the search box (Chapter 9.2)."),
        ("How do I schedule an appointment?",
         "Open Appointments → New Appointment, choose the patient, dentist, date, time, duration and type, then schedule (Chapter 10.1)."),
        ("How do I reschedule an appointment?",
         "Open the appointment → Edit, change the date/time/duration and save. The patient cannot be changed after booking."),
        ("How do I cancel an appointment?",
         "Open the appointment → Cancel and confirm. Completed appointments cannot be cancelled."),
        ("How do I create a prescription?",
         "Open the patient's record → Prescriptions tab, create the prescription, add medicines and save (Chapter 12)."),
        ("How do I print a prescription?",
         "Open the prescription and click Print; use the browser print dialog (Chapter 18)."),
        ("How do I download a prescription PDF?",
         "Open the prescription and click Download PDF; in the browser print dialog choose \u201cSave as PDF\u201d and save."),
        ("How do I create an invoice?",
         "Open Invoices → create, choose the patient and line items, then save as a Draft and Issue it (Chapter 14.2)."),
        ("How do I record a payment?",
         "Open Payments → record, choose the patient, amount and method, save, complete, then allocate to the invoice (Chapter 14.3)."),
        ("How do I print or download an invoice?",
         "Open the invoice and click Print or Download PDF; use the browser print dialog (Chapter 18)."),
        ("How do I upload an attachment?",
         "Open the patient's record → Attachments tab, upload the file (max 10 MB) and choose its type (Chapter 15)."),
        ("Why can't I see a particular module?",
         "Module visibility is controlled by your role. If the menu item is not shown (e.g. Users and Pending Approvals for non-administrators) or a screen says you lack permission, your role does not include it. Ask an administrator if you believe access should change."),
        ("Why can't I perform a particular action?",
         "Actions such as issuing invoices, approving refunds, deactivating patients or deleting documents are restricted to specific roles. Chapter 22 lists exactly who can do what."),
        ("What should I do if my session expires?",
         "Sign in again with your email and password. If you ticked \u201cRemember me\u201d, DensCare will keep you signed in on that computer."),
        ("What should I do if I cannot log in?",
         "Check the spelling of your email and password, confirm the account was approved and activated, and use \u201cForgot password\u201d to reset the password if needed. See Chapter 25."),
    ]
    for q, a in faqs:
        rich_para(doc, [(q, True, False, NAVY)], space_after=2)
        para(doc, a, space_after=8)

    page_break(doc)

    # ── Section 22 ──────────────────────────────────────────────────
    doc.add_heading("22.  Role-Based Quick Reference", level=1)

    para(doc,
         "The table below summarises what each role can do in the delivered "
         "system. Legend: " + CHECK + " = allowed, — = not available. The four "
         "doctor roles (General, Specialist, Consulting, Chief) share the same "
         "clinical permissions; the Chief Doctor additionally has administrator "
         "access, noted where it differs.")

    rows = [
        ("Dashboard", ["Y", "Y", "Y", "Y", "Y"]),
        ("Patients — view & search", ["Y", "Y", "Y", "Y", "N"]),
        ("Patients — register & edit", ["Y", "N", "N", "Y", "N"]),
        ("Patients — activate / deactivate", ["Y", "N", "N", "N", "N"]),
        ("Appointments — view / book / reschedule / cancel", ["Y", "Y", "Y", "Y", "N"]),
        ("Patient records — view & add content", ["Y", "Y", "Y", "Y", "N"]),
        ("Patient records — change status / finalize", ["Y", "Y", "Y", "N", "N"]),
        ("Patient records — delete (soft)", ["Y", "N", "N", "N", "N"]),
        ("Patient record audit trail — view", ["Y", "Y", "N", "N", "N"]),
        ("Prescriptions — create / print / download", ["Y", "Y", "Y", "Y", "N"]),
        ("Treatment plans — create & manage", ["Y", "Y", "Y", "Y", "N"]),
        ("Procedure catalogue — view", ["Y", "Y", "Y", "Y", "N"]),
        ("Procedure catalogue — manage", ["Y", "Y", "N", "N", "N"]),
        ("Doctors — view any doctor profile & schedule", ["Y", "N", "N", "Y", "N"]),
        ("Doctors — view own profile (doctor roles)", ["N", "Y", "Y", "N", "N"]),
        ("Doctors — create / edit / deactivate / schedules", ["Y", "Y", "N", "N", "N"]),
        ("Specializations — manage", ["Y", "Y", "N", "N", "N"]),
        ("Users — manage & approve registrations", ["Y", "Y", "N", "N", "N"]),
        ("Billing — view (dashboard, invoices, payments)", ["Y", "Y", "Y", "Y", "Y"]),
        ("Invoices — create & edit draft", ["Y", "Y", "Y", "Y", "Y"]),
        ("Invoices — issue / cancel", ["Y", "Y", "Y", "Y", "N"]),
        ("Invoices — delete draft", ["Y", "N", "N", "N", "N"]),
        ("Payments — record & edit", ["Y", "Y", "Y", "Y", "Y"]),
        ("Payments — complete / fail / void / allocate", ["Y", "Y", "Y", "Y", "N"]),
        ("Payments — delete", ["Y", "N", "N", "N", "N"]),
        ("Receipts — generate", ["Y", "Y", "Y", "Y", "Y"]),
        ("Receipts — regenerate", ["Y", "Y", "Y", "Y", "N"]),
        ("Refunds — create", ["Y", "Y", "Y", "Y", "Y"]),
        ("Refunds — approve / reject / complete", ["Y", "Y", "Y", "Y", "N"]),
        ("Credit notes — create", ["Y", "Y", "Y", "Y", "Y"]),
        ("Credit notes — issue / apply / void", ["Y", "Y", "Y", "Y", "N"]),
    ]
    headers = ["Feature / module", "Admin", "Chief Doctor", "Doctor", "Receptionist", "Dental Assistant"]
    table_rows = []
    for feat, marks in rows:
        cells = [feat]
        for m in marks:
            cells.append((CHECK if m == "Y" else (CROSS if m == "N" else m), m == "Y"))
        table_rows.append(cells)

    add_table(
        doc,
        headers,
        table_rows,
        col_widths=[Cm(6.4), Cm(2.0), Cm(2.1), Cm(2.1), Cm(2.3), Cm(2.1)],
        font_size=8.5,
    )

    callout(doc, "About this table",
            "These permissions were verified directly against the application's "
            "authorisation rules. If a screen shows \u201cyou do not have permission\u201d, "
            "the table above is the reason.")

    page_break(doc)

    # ── Section 23 ──────────────────────────────────────────────────
    doc.add_heading("23.  Client Training Checklist", level=1)

    para(doc, "Use these checklists while training the clinic staff. Tick each item "
              "when the trainee can do it without help.")

    doc.add_heading("23.1  Administrator training", level=2)
    for item in [
        "Sign in and sign out", "Dashboard and quick actions",
        "Create a user account / approve a pending registration and assign a role",
        "Change a user's role", "Activate and deactivate a user",
        "Create a doctor profile and assign specializations",
        "Set a doctor's weekly schedule and availability / leave",
        "Register a patient and open a patient profile",
        "Schedule, reschedule and cancel an appointment",
        "Open a patient record and review clinical content",
        "Create and manage a treatment plan",
        "Create, issue and cancel an invoice",
        "Record, complete and allocate a payment; generate a receipt",
        "Create, approve and complete a refund",
        "Create, issue and apply a credit note",
        "Upload and download an attachment",
        "Print / download an invoice and a prescription",
    ]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("23.2  Doctor training", level=2)
    for item in [
        "Sign in and sign out", "Dashboard — My Treatment Plans and today's appointments",
        "Search for a patient and open the profile",
        "Open an appointment and the patient's record",
        "Record a consultation (clinical details + medical history)",
        "Add and update diagnoses",
        "Create a prescription with medicines, dosage, frequency, duration and instructions",
        "Print / download a prescription",
        "Schedule a follow-up",
        "Create a treatment plan, add procedures and manage its status",
        "Create and issue an invoice for the patient",
        "Upload an attachment",
    ]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("23.3  Receptionist training", level=2)
    for item in [
        "Sign in and sign out", "Dashboard and quick actions",
        "Register a patient", "Search and open a patient",
        "Update patient information",
        "Schedule, reschedule and cancel an appointment",
        "Check today's appointments",
        "Create and issue an invoice",
        "Record, complete and allocate a payment",
        "Generate and print a receipt",
        "Create a refund (and know who approves it)",
        "Upload an attachment",
    ]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("23.4  Dental assistant training", level=2)
    for item in [
        "Sign in and sign out",
        "View the Billing Dashboard",
        "Create and edit a draft invoice",
        "Record a payment",
        "Generate a receipt",
        "Create a refund request and a draft credit note",
        "Understand which workflow actions require another role",
    ]:
        bullet(doc, item, bold_prefix="\u2610  ")

    page_break(doc)

    # ── Section 24 ──────────────────────────────────────────────────
    doc.add_heading("24.  Client Handover Checklist", level=1)

    para(doc, "Use this checklist when the application is handed over to the clinic.")

    doc.add_heading("24.1  Application", level=2)
    for item in ["Application URL provided and tested", "Login credentials provided securely",
                 "Roles explained to each user", "Navigation (sidebar, mobile drawer) explained",
                 "Mobile usage explained"]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("24.2  Administration", level=2)
    for item in ["Administrator account configured", "All staff users created",
                 "Roles assigned to every user", "Doctor accounts configured",
                 "Receptionist accounts configured"]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("24.3  Data", level=2)
    for item in ["Patient data verified", "Doctor data verified",
                 "Procedure catalogue verified (prices and categories)",
                 "Billing configuration verified (currencies, clinic settings)"]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("24.4  Training", level=2)
    for item in ["Administrator training completed", "Doctor training completed",
                 "Receptionist training completed", "Dental assistant training completed"]:
        bullet(doc, item, bold_prefix="\u2610  ")

    doc.add_heading("24.5  Documentation", level=2)
    for item in ["User manual delivered", "Credentials delivered securely",
                 "Support / contact information delivered"]:
        bullet(doc, item, bold_prefix="\u2610  ")

    page_break(doc)

    # ── Section 25 ──────────────────────────────────────────────────
    doc.add_heading("25.  Troubleshooting", level=1)

    doc.add_heading("25.1  \u201cI cannot log in\u201d", level=2)
    para(doc, "Possible reasons and what to do:")
    add_table(
        doc,
        ["Possible reason", "What to do"],
        [
            ["Incorrect email or password", "Check the spelling and try again. Use \u201cForgot password\u201d to reset the password."],
            ["Account not yet approved", "New registrations must be approved by an administrator before they can sign in. Ask an administrator to approve the account."],
            ["Account deactivated", "An administrator can reactivate the account."],
            ["Session/token issue", "Close the browser, reopen DensCare and sign in again."],
        ],
        col_widths=[Cm(6.0), Cm(10.0)],
    )

    doc.add_heading("25.2  \u201cI cannot see a module\u201d", level=2)
    para(doc,
         "What you see in the menu depends on your role. Two menu items (Users and "
         "Pending Approvals) are visible only to administrators and Chief Doctors. "
         "For other screens, if your role does not permit the module, DensCare "
         "shows a permission message when you open it. Chapter 22 explains who can "
         "access what. If you believe your access is wrong, ask an administrator.")

    doc.add_heading("25.3  \u201cA button is disabled / an action is unavailable\u201d", level=2)
    para(doc,
         "Actions are shown only when the document's state allows them. For "
         "example, an issued invoice cannot be edited, a completed appointment "
         "cannot be cancelled, and a refund can only be completed after it is "
         "approved. If an action is missing or greyed out, check the document's "
         "status and your role's permissions.")

    doc.add_heading("25.4  \u201cMy session expired\u201d", level=2)
    para(doc,
         "DensCare ends your session after a period of inactivity for security. "
         "Simply sign in again. On your own computer, tick \u201cRemember me\u201d to stay "
         "signed in longer.")

    doc.add_heading("25.5  \u201cThe PDF download opened a print dialog\u201d", level=2)
    para(doc,
         "This is how the current version works. In the browser's print dialog, "
         "choose \u201cSave as PDF\u201d as the destination (or print to a printer) and "
         "confirm. See Chapter 18.")

    doc.add_heading("25.6  \u201cThe patient list looks empty\u201d", level=2)
    para(doc,
         "Check the search box and the status filter — a search term or an "
         "Inactive filter may be hiding records. Deactivated patients are hidden "
         "from the default view.")

    doc.add_heading("25.7  Getting help", level=2)
    para(doc,
         "For anything not covered here, contact the DensCare support contact "
         "provided with your handover, and describe what you were doing when the "
         "problem appeared.")

    page_break(doc)

    # ── Section 26 ──────────────────────────────────────────────────
    doc.add_heading("26.  Future / Not Currently Available", level=1)

    para(doc,
         "The following items appear in DensCare's navigation or are commonly "
         "expected, but are NOT functional in the delivered version. They are "
         "listed here so nobody mistakes a placeholder for a working feature.")

    doc.add_heading("26.1  Shown in the navigation but disabled", level=2)
    add_table(
        doc,
        ["Item", "Status"],
        [
            ["Inventory", "Planned module — shown as a disabled menu item; not functional."],
            ["Laboratory", "Planned module — shown as a disabled menu item; not functional."],
            ["Reports", "Planned — shown as a disabled menu item; not functional."],
            ["Settings", "Planned — shown as a disabled menu item; not functional."],
        ],
        col_widths=[Cm(4.0), Cm(12.0)],
    )

    doc.add_heading("26.2  Not available in the current version", level=2)
    bullet(doc, "Patient portal — patients cannot log in to view their own records.")
    bullet(doc, "Calendar view — the Appointments list is the schedule view; there is no separate calendar screen.")
    bullet(doc, "Live dashboard statistics — the Overview cards and Recent Activity on the Dashboard show sample figures, not live data.")
    bullet(doc, "Standalone refund / credit-note lists — refunds and credit notes are created and viewed from the payment/invoice detail screens.")
    bullet(doc, "Advanced billing reports (revenue, cash flow, ageing, monthly/yearly reports) — not available; the Billing Dashboard is the reporting surface.")
    bullet(doc, "Editing an issued invoice's amounts — invoices become immutable once issued; corrections use cancellation or credit notes.")
    bullet(doc, "Patient record deletion by staff — only administrators can delete (soft-delete) patient records.")
    bullet(doc, "Automatic PDF file generation — PDF download uses the browser's print dialog (\u201cSave as PDF\u201d).")

    page_break(doc)

    # ── Closing ─────────────────────────────────────────────────────
    doc.add_heading("27.  Closing", level=1)
    para(doc,
         "This manual describes DensCare as delivered in Version 1.0. Every screen, "
         "action and permission has been checked against the application itself, so "
         "the clinic can rely on it during training and daily use.")
    para(doc,
         "Thank you for choosing DensCare. The DensCare Project Team wishes your "
         "clinic smooth, organised and successful operations.")
    para(doc, "", space_after=2)
    rich_para(doc, [("— End of manual —", True, True, SLATE)], align=WD_ALIGN_PARAGRAPH.CENTER)

    # ════════════════════════════════════════════════════════════════
    # Headers / footers
    # The document is a single section. `different_first_page_header_footer`
    # is already enabled, so the cover uses the (empty) first-page header and
    # every later page uses these default header/footer.
    # ════════════════════════════════════════════════════════════════
    for section in doc.sections:
        build_header(doc, section)
        build_footer(doc, section)

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


def add_module_table(doc, rows):
    """Purpose / Who uses / What it does / Features / Workflow / Actions / Related / Notes"""
    add_table(
        doc,
        ["", ""],
        rows,
        col_widths=[Cm(3.4), Cm(12.6)],
        font_size=9.5,
        first_col_bold=True,
    )
    doc.add_paragraph()


if __name__ == "__main__":
    build()
